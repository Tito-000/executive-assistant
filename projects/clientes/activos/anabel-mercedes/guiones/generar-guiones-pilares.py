"""Genera el documento Word con los 9 guiones de los 3 pilares de Anabel."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Márgenes
for section in doc.sections:
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# Estilos base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_title(text, level=1, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(22)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(15)
    elif level == 3:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    return p


def add_paragraph(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def add_meta(label, value):
    """Meta info: **Label:** value"""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(value)
    return p


def add_hooks(hooks):
    """3 opciones de hook"""
    p = doc.add_paragraph()
    r = p.add_run("Gancho hablado — 3 opciones:")
    r.bold = True
    for i, (hook, note) in enumerate(hooks, 1):
        pp = doc.add_paragraph(style='List Number')
        r = pp.add_run(f'"{hook}"')
        r.bold = True
        pp.add_run(f"  ({note})").italic = True


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"“{text}”")
    r.italic = True


def add_reentry(text):
    p = doc.add_paragraph()
    r = p.add_run(f"🔁 Re-entry: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run(f'"{text}"').italic = True


def add_direction(text):
    """Indicación de dirección escénica"""
    p = doc.add_paragraph()
    r = p.add_run(f"({text})")
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_screen_text(text):
    p = doc.add_paragraph()
    r = p.add_run("Texto en pantalla: ")
    r.bold = True
    p.add_run(text)


def add_horizontal_rule():
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


# ============== PORTADA ==============
add_title("GUIONES DE VIDEO", level=0, color=RGBColor(0x2C, 0x5F, 0x7F))
add_title("Anabel Mercedes — Los 3 Pilares de Contenido", level=2)
doc.add_paragraph()
add_paragraph("Fecha: 21 abril 2026", italic=True)
add_paragraph("Metodología: Skill /guion-viral (Victor + Ramiro fusionados)", italic=True)
add_paragraph("Contenido: 9 guiones (3 por pilar)", italic=True)
doc.add_paragraph()

add_paragraph("Estructura de cada guión:", bold=True)
doc.add_paragraph("• Idea ganadora + duración + plataforma + formato + target")
doc.add_paragraph("• Packaging (gancho visual + 3 hooks hablados + texto en pantalla)")
doc.add_paragraph("• Intro con re-entry")
doc.add_paragraph("• Cuerpo con escenas + open loops entre cada una")
doc.add_paragraph("• Outro (moral + CTA nativo)")

doc.add_paragraph()
add_paragraph("Regla de uso de los 3 hooks:", bold=True)
doc.add_paragraph("Anabel graba el cuerpo completo una sola vez y después graba las 3 variantes "
                  "de hook por separado (10 seg cada una). Con 3 tomas cortas obtiene 3 videos "
                  "distintos del mismo cuerpo — A/B/C testing orgánico sin triplicar el trabajo.")

doc.add_page_break()

# ============== PILAR 1 ==============
add_title("PILAR 1 — MÉTODO INTERMEDIARIO", level=1, color=RGBColor(0x2C, 0x5F, 0x7F))
add_paragraph("Los nuevos negocios ya no cargan producto, no emplean, no alquilan local. "
              "Son intermediarios entre empresas que ya tienen todo y clientes que lo necesitan.",
              italic=True)
doc.add_paragraph()

# --- VIDEO 1 ---
add_title("🎬 VIDEO 1 — Los negocios cambiaron", level=2)
add_meta("Idea ganadora", "El modelo de negocio del pasado ya no existe. Ahora se gana siendo intermediario, no dueño de todo.")
add_meta("Duración", "40-45 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "Talking Head + B-roll")
add_meta("Target", "Mujer que quiere emprender pero cree que necesita capital, local o inventario")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "Anabel en su casa, café en mano, mira a cámara con seguridad.")
add_hooks([
    ("Si tú todavía crees que para emprender necesitas un local, escucha esto…", "directo, contraintuitivo"),
    ("Los negocios cambiaron y nadie te lo dijo. Por eso tú todavía no arrancas.", "declaración + consecuencia"),
    ("POV: quieres emprender pero piensas que necesitas $10k pa' empezar.", "POV relatable"),
])
add_screen_text('"Los negocios cambiaron 👇"')
add_reentry("Y lo que te voy a decir va contra todo lo que tú piensas…")
doc.add_paragraph()

add_title("INTRO (4-10 seg)", level=3)
add_quote("Hace 20 años pa' tener un negocio tú necesitabas local, inventario, empleados, permisos, capital grueso. Hoy, nada de eso. Y si tú todavía piensas así, te estás quedando atrás.")
add_reentry("Porque el modelo ya no es ese. Mira…")
doc.add_paragraph()

add_title("CUERPO (10-35 seg)", level=3)

add_paragraph("Escena 1 — El modelo viejo (10-18 seg)", bold=True)
add_direction("B-roll de locales vacíos, cajas de inventario, facturas")
add_quote("Antes los negocios eran así — cargas producto, empleas gente, alquilas local, pagas luz, pagas impuestos. Y al final del mes te queda poquito.")
add_reentry("Pero eso no es ni lo peor…")
doc.add_paragraph()

add_paragraph("Escena 2 — El modelo nuevo (18-28 seg)", bold=True)
add_direction("Anabel hablando a cámara, energía alta")
add_quote("Los negocios de hoy son intermediarios. Conectan a una empresa que ya tiene todo — producto, logística, facturación — con un cliente que lo necesita. Tú no cargas nada. Tú conectas, capacitas y cobras.")
add_reentry("Y aquí viene lo que más le duele a la gente escuchar…")
doc.add_paragraph()

add_paragraph("Escena 3 — La verdad incómoda (28-35 seg)", bold=True)
add_quote("Lo que antes te costaba 10 años construir, hoy una persona lo arma en 6 meses con su celular. Si tú sigues pensando en 'poner un negocio', estás jugando el juego viejo.")
add_direction("Elementos narrativos usados: pensamiento, diálogo")
doc.add_paragraph()

add_title("OUTRO (35-45 seg)", level=3)
add_meta("Moral", '"Ya no se gana siendo dueño de todo. Se gana siendo el puente entre el que tiene y el que necesita."')
add_meta("CTA nativo", '"Si tú quieres ver cómo funciona este modelo nuevo por dentro, escríbeme MODELO y te lo explico."')
add_screen_text('"MODELO 👇"')
add_horizontal_rule()
doc.add_paragraph()

# --- VIDEO 2 ---
add_title("🎬 VIDEO 2 — Nadie te dice esto cuando decides emprender", level=2)
add_meta("Idea ganadora", "Cuando decides emprender, no te avisan que vas a perder amigos, que tu familia va a dudar de ti, y que el mayor obstáculo es mental.")
add_meta("Duración", "45-50 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "Momento Vulnerable / Selfie Casual")
add_meta("Target", "Mujer que está en los primeros meses de emprender y siente que nadie la entiende")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "Anabel en plano cerrado, luz natural, cero producción. Tono íntimo.")
add_hooks([
    ("Nadie te dice esto cuando decides emprender… y yo te lo voy a decir.", "promesa íntima"),
    ("Te voy a decir 3 cosas que nadie me dijo cuando yo empecé — y que me hubiera ahorrado mucho dolor.", "declaración de regalo"),
    ("Si tú acabas de arrancar un negocio y te sientes rara, no estás sola. Te explico lo que nadie te cuenta.", "identificación + invitación"),
])
add_screen_text('"Lo que nadie te cuenta 🤍"')
add_reentry("Y la tercera te va a doler…")
doc.add_paragraph()

add_title("INTRO (4-10 seg)", level=3)
add_quote("Cuando yo decidí emprender en serio, nadie me preparó pa' esto. Te lo voy a contar rápido pa' que tú no te sorprendas como me sorprendí yo.")
add_reentry("Y empieza por algo que no te esperas…")
doc.add_paragraph()

add_title("CUERPO (10-42 seg)", level=3)

add_paragraph("Escena 1 — Vas a perder amigos (10-20 seg)", bold=True)
add_quote("Uno — vas a perder amigos. Y no porque sean malos. Es que cuando tú empiezas a crecer, hay gente que no va pa' donde tú vas. Y eso duele, pero es parte.")
add_reentry("Y eso no es ni lo más difícil…")
doc.add_paragraph()

add_paragraph("Escena 2 — Tu familia va a dudar (20-30 seg)", bold=True)
add_quote("Dos — tu propia familia va a dudar de ti. Te van a decir 'consigue un trabajo normal', 'eso no funciona', 'déjate de inventar'. Y pues sí, duele, porque uno espera apoyo de ahí primero.")
add_reentry("Pero el tercero es el que te voltea la cara…")
doc.add_paragraph()

add_paragraph("Escena 3 — El mayor obstáculo eres tú (30-42 seg)", bold=True)
add_quote("Tres — y esta es la que nadie te dice — el mayor obstáculo no son ellos. Eres tú. Tu cabeza te va a tratar de convencer 50 veces al día de que te rindas. Y ganar ese diálogo interno es más difícil que conseguir clientes. Te lo juro.")
add_direction("Elementos narrativos usados: pensamiento, emociones, diálogo, ubicación")
doc.add_paragraph()

add_title("OUTRO (42-50 seg)", level=3)
add_meta("Moral", '"Emprender no es difícil por el dinero. Es difícil por lo que te toca desaprender."')
add_meta("CTA nativo", '"Si tú estás en esa etapa y necesitas que alguien te entienda, escríbeme EMPIEZO y hablamos."')
add_screen_text('"EMPIEZO 👇"')

doc.add_page_break()

# ============== PILAR 2 ==============
add_title("PILAR 2 — APALANCAMIENTO 1%", level=1, color=RGBColor(0x2C, 0x5F, 0x7F))
add_paragraph("El 99% quiere hacerlo todo solo. El 1% apalanca un sistema, duplica y deja de querer ser el mejor.",
              italic=True)
doc.add_paragraph()

# --- VIDEO 1 ---
add_title("🎬 VIDEO 1 — Si tú quieres hacerlo todo solo", level=2)
add_meta("Idea ganadora", "El emprendedor promedio fracasa porque quiere ser héroe. El 1% apalanca.")
add_meta("Duración", "45-50 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "Skit → Talking Head (Híbrido) — Anabel actúa dos versiones de sí misma: la caótica vs la tranquila, cierra hablando directo a cámara")
add_meta("Target", "Emprendedora que siente que 'nadie lo hace como ella'")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "SKIT: Anabel Versión A — pelo despeinado, papeles por todos lados, cargando cajas, café frío. Cámara rápida, caótica.")
add_hooks([
    ("Si tú también quieres hacerlo todo sola… te tengo una mala noticia.", "directo y retador"),
    ("POV: eres la emprendedora que hace todo sola y te preguntas por qué no creces.", "POV relatable"),
    ("Esto es lo que me dijo una mentora que me cambió el negocio: 'deja de querer ser la heroína'.", "storytelling con autoridad prestada"),
])
add_screen_text('"POV: eres la emprendedora-todo ✋"')
add_reentry("Y no es lo que tú estás pensando…")
doc.add_paragraph()

add_title("INTRO (4-12 seg)", level=3)
add_direction("Corte limpio — cambia de escena: Anabel Versión B, pelo arreglado, sentada tranquila, laptop, café caliente")
add_quote("Así pasaba yo antes. Creía que si yo no lo hacía, nadie lo iba a hacer bien. Hasta que me cayó la ficha de algo que me voló la cabeza…")
add_reentry("…y ahí entendí por qué yo iba a seguir estancada.")
doc.add_paragraph()

add_title("CUERPO (12-42 seg)", level=3)

add_paragraph("Escena 1 — El héroe se quema (12-22 seg)", bold=True)
add_direction("Corte: mezcla imágenes de la Versión A caótica mientras habla")
add_quote("Querer ser la mejor en todo te convierte en el cuello de botella de tu propio negocio. Tú vendes, tú das seguimiento, tú contestas, tú empacas — y a fin de mes estás agotada ganando lo mismo.")
add_reentry("Pero aquí viene lo que cambia el juego completo…")
doc.add_paragraph()

add_paragraph("Escena 2 — El 1% apalanca (22-33 seg)", bold=True)
add_direction("Vuelve a Anabel Versión B — talking head limpio, cercano")
add_quote("El 1% que factura en grande no trabaja más que tú. Trabaja distinto. Se monta arriba de un sistema que ya está probado, con gente que ya sabe, con herramientas que ya funcionan. No reinventan la rueda — la usan.")
add_reentry("Y eso es exactamente lo que a mí me cambió la vida…")
doc.add_paragraph()

add_paragraph("Escena 3 — Duplicar, no cargar (33-42 seg)", bold=True)
add_quote("Por eso a mí lo que me enamoró de este modelo fue eso — yo no estoy sola cargando. Yo duplico. Yo enseño lo que aprendí y la otra persona hace lo mismo. Así el ingreso se multiplica sin que tú te mates.")
add_direction("Elementos narrativos: pensamiento, emociones, diálogo, acciones")
doc.add_paragraph()

add_title("OUTRO (42-50 seg)", level=3)
add_meta("Moral", '"Emprender no es cargar más. Es cargar distinto."')
add_meta("CTA nativo", '"Si quieres ver este sistema por dentro, escríbeme APALANCA y te lo explico paso por paso."')
add_screen_text('"APALANCA 👇"')
add_horizontal_rule()
doc.add_paragraph()

# --- VIDEO 2 ---
add_title("🎬 VIDEO 2 — Deja de querer ser el mejor", level=2)
add_meta("Idea ganadora", "Creer que tienes que ser el #1 te mantiene chiquita.")
add_meta("Duración", "40-45 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "Confesión Cerrada — plano muy cerrado del rostro, luz natural, máxima intimidad")
add_meta("Target", "Mujer que se siente 'no lo suficientemente experta' para arrancar")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "Plano extremo cerrado del rostro de Anabel. Solo se le ven los ojos y parte de la boca. Luz natural de ventana. Sonrisa honesta que se vuelve seria.")
add_hooks([
    ("Te voy a ser honesta con algo que me costó años entender…", "confesional, íntimo"),
    ("El día que dejé de querer ser la mejor, mi negocio despegó. Y te explico por qué.", "declaración de resultado"),
    ("Si tú sientes que te falta 'estar más preparada' para arrancar, este video es para ti.", "identificación directa con la objeción"),
])
add_screen_text('"Confesión 🤍"')
add_reentry("…y cambió cómo yo veo este negocio.")
doc.add_paragraph()

add_title("INTRO (4-12 seg)", level=3)
add_direction("Mismo plano cerrado — cero cortes")
add_quote("Yo pasé mucho tiempo creyendo que para crecer yo tenía que ser la mejor. La más capacitada, la que más sabía, la que tuviera todas las respuestas. Y pues… estaba completamente equivocada.")
add_reentry("Porque mientras yo me enfocaba en eso, me estaba perdiendo lo más importante…")
doc.add_paragraph()

add_title("CUERPO (12-38 seg)", level=3)

add_paragraph("Escena 1 — La trampa del perfeccionismo (12-22 seg)", bold=True)
add_direction("Todavía mismo plano — tono sostenido")
add_quote("Cuando tú quieres ser la mejor, nunca arrancas. Siempre falta un curso más, una certificación más, un libro más. Y mientras tú te preparas, otra persona menos preparada que tú ya está facturando.")
add_reentry("Pero eso no es ni lo peor…")
doc.add_paragraph()

add_paragraph("Escena 2 — Ser la mejor te aísla (22-30 seg)", bold=True)
add_quote("Lo peor es que el que se cree el mejor no delega, no duplica, no construye equipo. Y sin equipo, tú no creces — tú te repites.")
add_reentry("Y ahí fue donde yo tomé una decisión…")
doc.add_paragraph()

add_paragraph("Escena 3 — El cambio de mentalidad (30-38 seg)", bold=True)
add_direction("Pequeña sonrisa — se suelta un poco")
add_quote("Yo dejé de querer ser la mejor y empecé a querer ser la mentora. Pasé de enfocarme en saber más… a enfocarme en que más gente supiera. Ahí fue que el ingreso empezó a moverse de verdad.")
add_direction("Elementos narrativos: pensamiento, emociones")
doc.add_paragraph()

add_title("OUTRO (38-45 seg)", level=3)
add_meta("Moral", '"Tu crecimiento no depende de cuánto sepas. Depende de cuántos aprendan contigo."')
add_meta("CTA nativo", '"Si sientes que te falta ser mejor antes de arrancar, escríbeme LISTA y hablamos."')
add_screen_text('"LISTA 👇"')
add_horizontal_rule()
doc.add_paragraph()

# --- VIDEO 3 ---
add_title("🎬 VIDEO 3 — La constancia no es lo que haces un día", level=2)
add_meta("Idea ganadora", "La gente confunde motivación con constancia. El 1% ejecuta sin ganas.")
add_meta("Duración", "35-40 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "Walk & Talk + Rant (Híbrido) — Anabel camina saliendo temprano mientras habla con convicción fuerte a cámara selfie")
add_meta("Target", "Mujer que arrancó con fuerza y se está apagando")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "Anabel caminando por la calle temprano en la mañana. Cámara selfie, energía fuerte, mira fijo al lente.")
add_hooks([
    ("La constancia no es lo que tú haces un día que tienes ganas. ESO no es constancia.", "opinión impopular / rant"),
    ("Opinión impopular: la motivación no existe. Y por eso tú no estás facturando.", "declaración fuerte + consecuencia"),
    ("¿Tú sabes por qué el 90% abandona en los primeros 6 meses? Te lo digo ahora.", "pregunta repentina + promesa"),
])
add_screen_text('"Se van a molestar 🔥"')
add_reentry("Es lo que haces el día que NO tienes ni una gota de ganas…")
doc.add_paragraph()

add_title("INTRO (4-10 seg)", level=3)
add_direction("Sigue caminando — ahora sube unos escalones")
add_quote("Y esto es lo que separa a la que factura de la que abandona. No es talento. No es suerte. Es esto…")
add_reentry("…y lo que voy a decir ahora le molesta a mucha gente.")
doc.add_paragraph()

add_title("CUERPO (10-34 seg)", level=3)

add_paragraph("Escena 1 — La motivación miente (10-20 seg)", bold=True)
add_direction("Se detiene en una esquina, mira fijo a cámara, energía alta")
add_quote("La motivación es una emoción. Llega, te prende, y a los 3 días se va. Si tú construyes tu negocio arriba de motivación, te va a ir mal — porque los días que no estés motivada, no vas a hacer NADA.")
add_reentry("Entonces, ¿cuál es la diferencia entre la que llega y la que no?")
doc.add_paragraph()

add_paragraph("Escena 2 — La constancia es sistema (20-34 seg)", bold=True)
add_direction("Retoma la caminata — tono se suaviza, más resuelto")
add_quote("La que llega tiene un sistema que no depende de cómo se siente. Tiene una hora para publicar, una hora para contactar, una hora para capacitarse — y lo hace llueva, truene, o le dé flojera. Punto. El 1% no es más disciplinado que tú. Solo tiene el sistema que a ti te falta.")
add_direction("Elementos narrativos: ubicación, acciones, diálogo")
doc.add_paragraph()

add_title("OUTRO (34-40 seg)", level=3)
add_direction("Llega a un café, se sienta, cámara fija")
add_meta("Moral", '"La motivación arranca. El sistema sostiene. La constancia paga."')
add_meta("CTA nativo", '"Si quieres saber cuál es el sistema que yo sigo todos los días, escríbeme SISTEMA."')
add_screen_text('"SISTEMA 👇"')
add_horizontal_rule()
doc.add_paragraph()

# --- VIDEO 4 ---
add_title("🎬 VIDEO 4 — Vive tu proceso", level=2)
add_meta("Idea ganadora", "La gente se frustra comparando su mes 2 con el año 5 de otra persona.")
add_meta("Duración", "40-45 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "Day in the Life + B-roll + Voiceover (Híbrido) — secuencia de momentos reales del día de Anabel con voiceover reflexivo")
add_meta("Target", "Mujer frustrada porque 'no ve resultados rápido'")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "B-roll slow motion — Anabel mirando por la ventana con café en mano, luz de la mañana entrando. Música suave, cinemática.")
add_hooks([
    ("Si llevas 3 meses y te quieres rendir, escucha esto…", "identificación con dolor específico"),
    ("Deja de compararte con gente que lleva 5 años. Te estás autosaboteando.", "declaración directa, casi reproche"),
    ("Para ti, que estás cansada y sientes que no estás avanzando — este mensaje es tuyo.", "tono cálido, abrazo"),
])
add_screen_text('"Para ti, que estás cansada 🤍"')
add_reentry("…porque esto me hubiera ahorrado mucho dolor.")
doc.add_paragraph()

add_title("INTRO (4-12 seg)", level=3)
add_direction("B-roll: Anabel cerrando los ojos, respirando, poniendo el celular en la mesa")
add_quote("Yo también llegué a ese punto. Mirando a gente que llevaba 2, 3, 5 años en esto, comparando mis resultados con los de ellos. Y pues… me estaba autosaboteando.")
add_reentry("Y esto es lo que nadie te dice cuando arrancas…")
doc.add_paragraph()

add_title("CUERPO (12-38 seg)", level=3)

add_paragraph("Escena 1 — La comparación es veneno (12-22 seg)", bold=True)
add_direction("B-roll: scroll de Instagram, mirando historias de otras emprendedoras con logros")
add_quote("El problema no es que tú no estés creciendo. El problema es que estás midiendo tu mes 2 con el año 5 de otra persona. Y no hay manera de que eso se sienta bien. Siempre vas a perder en esa comparación.")
add_reentry("Pero aquí está lo que cambia todo…")
doc.add_paragraph()

add_paragraph("Escena 2 — Cada proceso es único (22-32 seg)", bold=True)
add_direction("B-roll: Anabel caminando con sus hijos, luego capacitándose, luego grabando contenido")
add_quote("Tu proceso es TUYO. Tu ritmo, tus ventanas de tiempo, tu punto de partida. Hay gente que lleva 5 años y arrancó con ventajas que tú no tienes. Y hay gente que en 6 meses está donde otros en 3 años. Nadie corre la misma carrera.")
add_reentry("Y cuando yo entendí esto…")
doc.add_paragraph()

add_paragraph("Escena 3 — El proceso ES el negocio (32-38 seg)", bold=True)
add_direction("B-roll cinemático: Anabel al final del día, apagando la laptop, sonriendo satisfecha, vaso de Immunocal en la mesa")
add_quote("…paré de compararme y empecé a disfrutar mi camino. Cada 'no', cada lead que se caía, cada día de dudas — todo eso era parte de convertirme en la persona que hoy factura.")
add_direction("Elementos narrativos: ubicación, acciones, pensamiento, emociones")
doc.add_paragraph()

add_title("OUTRO (38-45 seg)", level=3)
add_meta("Moral", '"El proceso no es lo que te trae los resultados. El proceso ES el resultado."')
add_meta("CTA nativo", '"Si estás en ese punto de dudas, escríbeme PROCESO y hablamos."')
add_screen_text('"PROCESO 👇"')

doc.add_page_break()

# ============== PILAR 3 ==============
add_title("PILAR 3 — SALUD DE ALTO PERFORMANCE", level=1, color=RGBColor(0x2C, 0x5F, 0x7F))
add_paragraph("La salud no es un gasto, es una inversión. Diferencia entre mantener y reparar. Ritual, no prisa.",
              italic=True)
doc.add_paragraph()

# --- VIDEO 1 ---
add_title("🎬 VIDEO 1 — Diferencia Immunocal Azul vs Platinum", level=2)
add_meta("Idea ganadora", "La mayoría compra el Azul sin saber que el Platinum resuelve problemas distintos.")
add_meta("Duración", "45-50 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "Comparación + Paso a Paso Visual (Híbrido) — muestra los dos productos lado a lado")
add_meta("Target", "Persona que ya conoce Immunocal pero duda cuál comprar")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "Anabel con un sobre Azul en una mano y un sobre Platinum en la otra, mirándolos como si estuviera decidiendo.")
add_hooks([
    ("Si tú estás por comprar Immunocal y no sabes cuál, para este video…", "identificación con el momento de decisión"),
    ("Azul o Platinum — el 80% elige mal y pierde meses de resultados.", "dato + consecuencia, crea urgencia"),
    ("POV: estás frente a los dos sobres de Immunocal y no tienes idea cuál agarrar.", "POV relatable"),
])
add_screen_text('"AZUL vs PLATINUM 🧬"')
add_reentry("…porque elegir mal te puede hacer perder meses.")
doc.add_paragraph()

add_title("INTRO (4-10 seg)", level=3)
add_quote("A mí me preguntan esto como 10 veces al día. Y pues, no es que uno sea mejor que el otro — es que son PARA cosas distintas. Te lo explico facilito.")
add_reentry("Y lo que te voy a decir ahora lo cambia todo…")
doc.add_paragraph()

add_title("CUERPO (10-42 seg)", level=3)

add_paragraph("Escena 1 — El Azul: mantenimiento (10-22 seg)", bold=True)
add_direction("Muestra el sobre azul en pantalla, primer plano")
add_quote("El Azul es tu base. Es para el día a día — mantener el sistema inmune fuerte, que no te enfermes cada dos semanas, que tengas energía estable. Si tú estás sana y quieres prevenir, este es el tuyo.")
add_reentry("Pero si tu cuerpo ya está pidiendo auxilio, necesitas otra cosa…")
doc.add_paragraph()

add_paragraph("Escena 2 — El Platinum: carga pesada (22-36 seg)", bold=True)
add_direction("Cambia al sobre platinum, lo levanta")
add_quote("El Platinum es más potente. Trae los mismos precursores del Azul pero con RMF y fórmula más cargada. Este es para cuando tu cuerpo está bajo estrés real — mucha carga laboral, entrenamiento fuerte, o estás saliendo de algo serio. Aquí no estás manteniendo, estás reparando.")
add_reentry("Entonces, ¿cuál te toca a ti?")
doc.add_paragraph()

add_paragraph("Escena 3 — Cómo decidir (36-42 seg)", bold=True)
add_quote("Pregúntate esto: ¿yo estoy bien y quiero seguir bien? Azul. ¿Mi cuerpo está gritando y necesito apoyo fuerte? Platinum. Así de simple.")
add_direction("Elementos narrativos: diálogo, pensamiento")
doc.add_paragraph()

add_title("OUTRO (42-50 seg)", level=3)
add_meta("Moral", '"No es cuál es más caro. Es cuál es para TU momento de vida."')
add_meta("CTA nativo", '"Si todavía tienes dudas de cuál te toca, escríbeme IMMUNO y lo vemos juntas según tu caso."')
add_screen_text('"IMMUNO 👇"')
add_horizontal_rule()
doc.add_paragraph()

# --- VIDEO 2 ---
add_title("🎬 VIDEO 2 — Invierte en tu salud para no gastar en enfermedad", level=2)
add_meta("Idea ganadora", "La gente dice que la salud es cara — pero la enfermedad es infinitamente más cara.")
add_meta("Duración", "40-45 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "Rant / Desahogo + Opinión Impopular (Híbrido) — energía alta, convicción")
add_meta("Target", "Persona que dice 'no tengo dinero para Immunocal'")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "Anabel mirando fijo a cámara, medio seria, señalando con el dedo. Cero B-roll — pura energía.")
add_hooks([
    ("Me cansé de escuchar que Immunocal es caro. Déjame decirte algo…", "rant con autoridad"),
    ("Opinión impopular: no es que Immunocal sea caro — es que todavía no has visto la factura de enfermarte.", "declaración + reencuadre"),
    ("¿Tú dices que no tienes dinero pa' cuidarte? Entonces tampoco vas a tener pa' curarte.", "pregunta retórica + consecuencia"),
])
add_screen_text('"Opinión impopular 🔥"')
add_reentry("…y esto a mucha gente no le va a gustar.")
doc.add_paragraph()

add_title("INTRO (4-10 seg)", level=3)
add_quote("Todo el mundo dice 'eso está caro, no puedo pagarlo'. Y yo siempre les pregunto lo mismo, una sola cosa…")
add_reentry("…una pregunta que los deja callados.")
doc.add_paragraph()

add_title("CUERPO (10-36 seg)", level=3)

add_paragraph("Escena 1 — La pregunta que duele (10-20 seg)", bold=True)
add_quote("¿Tú sabes cuánto cuesta UNA noche en un hospital? ¿Una consulta con un especialista? ¿Los medicamentos cuando ya el cuerpo te falló? Multiplica eso por lo que Dios quiera y verás lo que es caro de verdad.")
add_reentry("Pero eso no es ni la mitad del problema…")
doc.add_paragraph()

add_paragraph("Escena 2 — El costo invisible (20-30 seg)", bold=True)
add_quote("Lo peor no es el dinero que gastas cuando te enfermas. Lo peor son los días que no trabajas, los planes que cancelas, la energía que no tienes pa' tus hijos, el humor de perro con el que andas. Eso no tiene precio — y lo pagas igual.")
add_reentry("Y aquí viene la parte incómoda…")
doc.add_paragraph()

add_paragraph("Escena 3 — La verdad (30-36 seg)", bold=True)
add_quote("Invertir en tu salud AHORA que estás bien es la inversión más barata que vas a hacer. El problema no es que Immunocal sea caro — el problema es que todavía no has visto la factura de no cuidarte.")
add_direction("Elementos narrativos: diálogo, emociones, pensamiento")
doc.add_paragraph()

add_title("OUTRO (36-45 seg)", level=3)
add_meta("Moral", '"La salud no es un gasto. Es la única inversión que siempre paga."')
add_meta("CTA nativo", '"Si tú quieres empezar a cuidarte sin quebrarte, escríbeme SALUD y te armo un plan según tu bolsillo."')
add_screen_text('"SALUD 👇"')
add_horizontal_rule()
doc.add_paragraph()

# --- VIDEO 3 ---
add_title("🎬 VIDEO 3 — Comparte y shake (barriendo el Immunocal)", level=2)
add_meta("Idea ganadora", "Mostrar el ritual real de preparación — no explicar beneficios, ENSEÑAR el momento.")
add_meta("Duración", "30-35 seg")
add_meta("Plataforma", "Reel")
add_meta("Formato", "POV mañana + Caption-Driven (Híbrido) — POV de ella preparando el shake, con texto grande para silent scrollers")
add_meta("Target", "Gente que ya compró Immunocal pero no lo toma bien + curiosos")
doc.add_paragraph()

add_title("PACKAGING (0-4 seg)", level=3)
add_meta("Gancho visual", "POV — manos de Anabel abriendo el sobre de Immunocal, cámara cerca. Música trending de fondo.")
add_hooks([
    ("POV: así tomo Immunocal yo todas las mañanas…", "POV ritual, casual"),
    ("Si tú tomas Immunocal y no haces ESTO al final, estás botando tu dinero.", "advertencia + promesa de truco"),
    ("La mayoría toma Immunocal mal. Y no se dan cuenta. Mira el truco…", "declaración + invitación"),
])
add_screen_text('"EL RITUAL QUE NO ME SALTO 🥤"')
add_reentry("(texto en pantalla) Espera el truco al final 👀")
doc.add_paragraph()

add_title("INTRO visual (4-8 seg)", level=3)
add_direction("Secuencia rápida de cortes — sin voz, música dominando: vaso de vidrio con agua fría / cuchara / sobre de Immunocal abriéndose")
add_screen_text('"Paso 1: agua fría, nunca caliente ❌🔥"')
add_reentry("(texto) Paso 2 es donde casi todos fallan…")
doc.add_paragraph()

add_title("CUERPO (8-30 seg)", level=3)

add_paragraph("Escena 1 — Cómo se echa (8-15 seg)", bold=True)
add_direction("Voz en off, casual")
add_quote("Mira, agua fría porque el calor mata las proteínas. Echas el sobre ASÍ — despacio, no de golpe — porque si lo tiras de una se te arrachima todo.")
add_screen_text('"Despacio, no de golpe ✋"')
add_reentry("(texto) Ahora el paso que nadie te enseña…")
doc.add_paragraph()

add_paragraph("Escena 2 — El shake correcto (15-22 seg)", bold=True)
add_direction("Muestra cómo mueve — NO batidora, cuchara circular suave")
add_quote("Y aquí está el truco — no uses batidora, no uses licuadora. Revuelves con cuchara, suavecito, en círculos. La batidora te rompe la proteína y pierdes todo el beneficio.")
add_screen_text('"Cuchara > Batidora siempre"')
add_reentry("(texto) Y antes de tomártelo…")
doc.add_paragraph()

add_paragraph("Escena 3 — Barrer (22-30 seg)", bold=True)
add_direction("Voz + visual: ella pasa la cuchara por el fondo del vaso")
add_quote("Y lo más importante — barre el fondo. Ahí es donde se queda todo el producto bueno. Si tú no barres, te estás tomando agua con gusto. Lo que vale está abajo.")
add_screen_text('"BARRE EL FONDO 🥄"')
add_direction("Elementos narrativos: ubicación, acciones, diálogo")
doc.add_paragraph()

add_title("OUTRO (30-35 seg)", level=3)
add_meta("Moral", '"Un buen producto mal tomado = dinero botado. Es ritual, no es prisa."')
add_meta("CTA nativo", '"Si quieres mi rutina completa de la mañana con Immunocal, escríbeme RITUAL y te la paso."')
add_screen_text('"RITUAL 👇"')

# ============== CIERRE ==============
doc.add_page_break()
add_title("📊 RESUMEN GENERAL", level=1, color=RGBColor(0x2C, 0x5F, 0x7F))
doc.add_paragraph()
add_paragraph("9 guiones · 3 pilares · 27 variantes de hook (A/B/C)", bold=True)
doc.add_paragraph()

add_paragraph("Pilar 1 — Método Intermediario (2 videos):", bold=True)
doc.add_paragraph("• Video 1: Los negocios cambiaron (Talking Head + B-roll)")
doc.add_paragraph("• Video 2: Nadie te dice esto cuando decides emprender (Momento Vulnerable)")
doc.add_paragraph()

add_paragraph("Pilar 2 — Apalancamiento 1% (4 videos):", bold=True)
doc.add_paragraph("• Video 1: Si tú quieres hacerlo todo solo (Skit → Talking Head)")
doc.add_paragraph("• Video 2: Deja de querer ser el mejor (Confesión Cerrada)")
doc.add_paragraph("• Video 3: La constancia no es lo que haces un día (Walk & Talk + Rant)")
doc.add_paragraph("• Video 4: Vive tu proceso (Day in the Life + B-roll VO)")
doc.add_paragraph()

add_paragraph("Pilar 3 — Salud de Alto Performance (3 videos):", bold=True)
doc.add_paragraph("• Video 1: Diferencia Immunocal Azul vs Platinum (Comparación + Paso a Paso)")
doc.add_paragraph("• Video 2: Invierte en tu salud (Rant + Opinión Impopular)")
doc.add_paragraph("• Video 3: Comparte y shake (POV mañana + Caption-Driven)")
doc.add_paragraph()

add_paragraph("Reglas de producción:", bold=True)
doc.add_paragraph("• Cada video se graba UNA vez en cuerpo, y después se graban las 3 variantes de hook")
doc.add_paragraph("• Respetar el formato indicado — no todos son talking head")
doc.add_paragraph("• Subtítulos siempre (80%+ ven Reels sin sonido)")
doc.add_paragraph("• CTA con keyword al final de cada Reel")
doc.add_paragraph("• Tono: dominicano natural, cercano, anti-hype")

# Guardar
output_path = "/Users/martinmercedes/Desktop/Executive assistant 2/projects/clientes/activos/anabel-mercedes/guiones/guiones-3-pilares-anabel.docx"
doc.save(output_path)
print(f"✅ Documento generado: {output_path}")
