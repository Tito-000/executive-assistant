from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# === CONFIGURACIÓN GENERAL ===
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(text, bold=False, align=None, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if align:
        p.alignment = align
    return p

def add_image_if_exists(path, width_inches=6):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

# ============================================================
# PORTADA
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIBE')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Mercadeo')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Maestría de Mercadeo')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistemas de Información Gerencial')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Consultoría de Mercadeo Digital para Farmacias Carol, S.R.L.\nAplicación de Plataformas Digitales y Sistemas de Gestión de Contenido')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Martin Mercedes\nMatrícula: [TU MATRÍCULA]')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Profesor: Eliezer Figueroa')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('9 de abril de 2026')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# ÍNDICE
# ============================================================
add_heading('Índice', level=1)

indice_items = [
    ('1. Introducción', '3'),
    ('', ''),
    ('MÓDULO 5: Diagramación de la Arquitectura Web', '4'),
    ('   5.1 Marco conceptual', '4'),
    ('   5.2 Diagnóstico de la arquitectura actual', '4'),
    ('   5.3 Propuesta de nueva arquitectura web', '5'),
    ('   5.4 Wireframes de las pantallas clave', '6'),
    ('   5.5 SEO y arquitectura web', '7'),
    ('', ''),
    ('MÓDULO 6: Manejadores de Contenido para Sitios Web (CMS)', '8'),
    ('   6.1 Marco conceptual', '8'),
    ('   6.2 Diagnóstico del CMS actual de Farmacia Carol', '8'),
    ('   6.3 Evaluación de opciones de CMS', '9'),
    ('   6.4 Recomendación final: WordPress + WooCommerce', '10'),
    ('   6.5 Arquitectura propuesta del sitio (Mockups HD)', '10'),
    ('   6.6 Plan de implementación y presupuesto', '11'),
    ('', ''),
    ('Conclusión', '12'),
    ('Bibliografía', '13'),
    ('Declaración de uso de IA', '14'),
]

for item, pag in indice_items:
    if item == '':
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(5.5))
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if not item.startswith('   '):
        run.bold = True
    p.add_run('\t' + pag).font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# INTRODUCCIÓN
# ============================================================
add_heading('1. Introducción', level=1)

intro = (
    "El presente trabajo tiene como finalidad aplicar los conocimientos adquiridos en la clase "
    "al analizar un negocio local en la República Dominicana. En este caso, elegimos a Farmacia Carol, S.R.L., "
    "una de las cadenas farmacéuticas más reconocidas de la República Dominicana."
)
add_paragraph(intro)

contexto = (
    "Para dar contexto de la empresa, Farmacia Carol fue fundada en el año 1977 en Santiago de los Caballeros "
    "por Don José Manuel Cabrera, conocido casualmente como Papá Carol. Esta empezó con una sola farmacia y "
    "con el pasar de los años hoy en día se ha convertido en una de las principales farmacéuticas más grandes "
    "de todo el país, con más de 60 sucursales activas distribuidas en todo el territorio nacional. Su modelo "
    "de negocio funciona con la venta de medicamentos y productos de cuidado personal en tienda física, y "
    "servicio de delivery a domicilio disponible las 24 horas en varias sucursales."
)
add_paragraph(contexto)

desafios = (
    "A pesar de su crecimiento y reconocimiento de marca, Farmacia Carol enfrenta diferentes desafíos en su "
    "presencia digital. Aunque la empresa cuenta con sitio web y presencia en las redes sociales, estas presentan "
    "claras oportunidades de mejora."
)
add_paragraph(desafios)

rol = (
    "En este trabajo se asume el rol de consultor de mercadeo digital contratado por la empresa para "
    "diagnosticar su situación actual y proponer soluciones concretas. Se seleccionaron los siguientes módulos:"
)
add_paragraph(rol)

modulo5_intro = (
    "Módulo 5 — Diagramación de la Arquitectura Web: aquí se analizará cómo se estructura el sitio web actual "
    "de Farmacia Carol y se propone una nueva arquitectura que mejore la navegación, la experiencia de usuario "
    "y el posicionamiento."
)
add_paragraph(modulo5_intro)

modulo6_intro = (
    "Módulo 6 — Manejadores de Contenido para Sitios Web (CMS): se evaluarán los diferentes CMS y se "
    "recomendará la plataforma más adecuada para que la empresa pueda gestionar su presencia digital de forma "
    "autónoma, eficiente y escalable."
)
add_paragraph(modulo6_intro)

competencia = (
    "Farmacia Carol es una de las marcas más fuertes en su sector, compitiendo con GBC y Los Hidalgos, lo "
    "cual representa un riesgo competitivo pero también una oportunidad de crecimiento."
)
add_paragraph(competencia)

doc.add_page_break()

# ============================================================
# MÓDULO 5
# ============================================================
add_heading('MÓDULO 5: Diagramación de la Arquitectura Web', level=1)

# --- 5.1 ---
add_heading('5.1 Marco conceptual', level=2)

p51_1 = (
    "La arquitectura web es básicamente la forma en la que se organiza, estructura y conecta la información "
    "dentro de un sitio web. Este viene siendo el plano o el mapa que define las diferentes partes de la página "
    "web y cómo los motores de búsqueda como Google interpretan y clasifican la información del sitio."
)
add_paragraph(p51_1)

p51_2 = (
    "Una diagramación de arquitectura web es la representación visual de cómo quedará la página. Esta tiene la "
    "finalidad de mostrar todas y cada una de las páginas del sitio, sus niveles de jerarquía y la conexión "
    "entre ellas. Hacer una diagramación es un paso muy importante porque permite planificar la experiencia "
    "del usuario antes de tirar una sola línea de código."
)
add_paragraph(p51_2)

p51_3 = (
    "En el caso del presente trabajo, realizaremos un wireframe para diagramar la arquitectura web de Farmacia "
    "Carol. Este es un boceto tipo esquema que representa la estructura visual de una página web sin incluir "
    "colores, tipografías finales ni imágenes reales. Esta cumple con la función de validar la usabilidad y "
    "la lógica de navegación antes de invertir en el desarrollo."
)
add_paragraph(p51_3)

# --- 5.2 ---
add_heading('5.2 Diagnóstico de la arquitectura actual', level=2)

p52_1 = (
    "Cuando analizamos la página de Farmacia Carol (www.farmaciacarol.com) se identifican dos problemas "
    "principales. No es una sola plataforma, sino dos sistemas que conviven bajo la misma marca."
)
add_paragraph(p52_1)

p52_2 = (
    "El portal corporativo corre sobre Microsoft SharePoint 2013 y aloja las secciones institucionales. La "
    "tienda en línea corre sobre VevoCart en un subdominio aparte, con su propio header, menú, login y "
    "carrito. Al hacer clic en \"Tienda\" el usuario salta a otra interfaz completamente distinta. A esto se "
    "suman tres subdominios más (InfoCarol, Caroleal, MommyClub-Caroleal), sumando cinco sitios inconexos "
    "sin sesión compartida."
)
add_paragraph(p52_2)

p52_3 = (
    "También pudimos identificar estos otros errores dentro de la plataforma web de Farmacia Carol:"
)
add_paragraph(p52_3)

errores = [
    "Catálogo de 4,856 artículos sin filtros avanzados, haciendo difícil la búsqueda.",
    "Sus páginas de productos no muestran mucha información: solo nombre, imagen del producto, precio y botón de añadir al carrito.",
    "Contenido desactualizado: la página principal muestra noticias del 2023, cometiendo un error crítico en el SEO.",
    "Enlaces muertos: varios enlaces de imágenes de productos no llevan a ninguna parte.",
    "Sin WhatsApp ni app móvil visibles, en un país donde el WhatsApp es el canal de comunicación dominante.",
    "Los activos \"Buscar Doctor\" y \"Consulta al Farmacéutico\" están siendo desaprovechados, ya que se encuentran en la parte inferior sin llamados a la acción fuertes.",
    "Búsqueda desactualizada y rota, sin filtros fáciles de entender.",
    "La sección de localización de las 60 sucursales se ve vieja y obsoleta.",
]

for e in errores:
    p = doc.add_paragraph(e, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# --- 5.3 ---
add_heading('5.3 Propuesta de nueva arquitectura web', level=2)

p53_1 = (
    "La nueva arquitectura para Farmacia Carol resolverá el problema de raíz en un solo sitio web. Se propone "
    "una arquitectura jerárquica de 3 niveles sobre una plataforma única (WordPress + WooCommerce o Shopify "
    "Plus) que eliminará los subdominios dispersos."
)
add_paragraph(p53_1)

p53_2 = (
    "(Diagrama generado con Excalidraw a partir de un prompt a IA. Ver abajo el prompt utilizado y la imagen resultante.)"
)
p = doc.add_paragraph()
run = p.add_run(p53_2)
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

add_paragraph('Prompt utilizado (herramienta: Claude + skill excalidraw-diagram):', bold=True)

prompt_excalidraw = (
    "\"Crear diagrama jerárquico del mapa de sitio propuesto para Farmacia Carol. Estructura de 3 niveles "
    "partiendo del nodo raíz 'FarmaciaCarol.com (Home)'. Nivel 1 (7 nodos): Productos, Sucursales, Servicios, "
    "Caroleal, Blog/InfoCarol, Sobre Nosotros, Contacto. Nivel 2 (subcategorías): Productos → Medicamentos, "
    "Salud y Bienestar, Dermocosmética, Cuidado Personal, Bebé y Mamá, Minimarket; Sucursales → Buscador con "
    "Mapa, Sucursales 24h, Sucursales con Delivery; Servicios → Delivery a Domicilio, Buscar Doctor, Consulta "
    "al Farmacéutico, AyudaMed. Estilo: organigrama top-down, cajas con bordes redondeados, nodo Home "
    "destacado en azul, nivel 1 en verde, nivel 2 en gris claro.\""
)
p = doc.add_paragraph()
run = p.add_run(prompt_excalidraw)
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.italic = True

# Insertar imagen del mapa de sitio si existe
mapa_path = '/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/wireframes/mapa-sitio-farmacia-carol.png'
if add_image_if_exists(mapa_path, width_inches=6.5):
    add_caption('Figura 1. Mapa de sitio propuesto para Farmacia Carol.')
else:
    add_paragraph('[INSERTAR IMAGEN: mapa-sitio-farmacia-carol.png]', bold=True)
    add_caption('Figura 1. Mapa de sitio propuesto para Farmacia Carol.')

# --- 5.4 ---
add_heading('5.4 Wireframes de las pantallas clave', level=2)

p54_1 = (
    "Diseñamos 5 wireframes que resuelven los puntos críticos mencionados en el punto 5.2:"
)
add_paragraph(p54_1)

wireframes_desc = [
    "Home: unifica los 5 subdominios en un solo sitio con navegación clara, hero, 6 categorías, más vendidos y localizador.",
    "Ficha de producto: agrega disponibilidad por sucursal, badge de receta, productos relacionados y reviews — hoy no existen.",
    "Categoría con filtros: filtros facetados (marca, precio, forma farmacéutica, receta) para navegar los 4,856 productos.",
    "Sucursales: mapa interactivo con filtros por ciudad, servicio 24h y delivery — reemplaza el listado plano actual.",
    "Checkout: flujo de 3 pasos con sticky order summary y descuento Caroleal integrado.",
]

for w in wireframes_desc:
    p = doc.add_paragraph(w, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

p54_2 = (
    "(Wireframes generados con IA — herramienta: nano-banana-images / Kie.ai con Gemini. Ver prompts completos en la Declaración de Uso de IA al final del documento.)"
)
p = doc.add_paragraph()
run = p.add_run(p54_2)
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

# Insertar los 5 wireframes
wireframe_files = [
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/wireframes/wf1-home.jpg', 'Figura 2. Wireframe — Home.'),
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/wireframes/wf2-product.jpg', 'Figura 3. Wireframe — Ficha de producto.'),
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/wireframes/wf3-category.jpg', 'Figura 4. Wireframe — Categoría con filtros.'),
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/wireframes/wf4-sucursales.jpg', 'Figura 5. Wireframe — Sucursales.'),
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/wireframes/wf5-checkout.jpg', 'Figura 6. Wireframe — Checkout.'),
]

for path, caption in wireframe_files:
    if add_image_if_exists(path, width_inches=5):
        add_caption(caption)
    else:
        add_paragraph(f'[INSERTAR IMAGEN: {os.path.basename(path)}]', bold=True)
        add_caption(caption)

# --- 5.5 ---
add_heading('5.5 SEO y arquitectura web', level=2)

p55 = (
    "Una página web bien organizada es una página que Google puede entender. Hoy Farmacia Carol tiene errores "
    "básicos graves: el título de la página está vacío, no tiene mapa del sitio y las direcciones web son "
    "números sin sentido. La nueva arquitectura arregla todo eso con URLs claras que le indican a Google qué "
    "es cada cosa de manera correcta. El objetivo es aumentar la cantidad de páginas indexadas de Farmacia "
    "Carol y convertirla en una verdadera tienda online, a la altura de su trayectoria en el mercado de más "
    "de 10 años."
)
add_paragraph(p55)

doc.add_page_break()

# ============================================================
# MÓDULO 6
# ============================================================
add_heading('MÓDULO 6: Manejadores de Contenido para Sitios Web (CMS)', level=1)

# --- 6.1 ---
add_heading('6.1 Marco conceptual', level=2)

p61 = (
    "Los Sistemas de Gestión de Contenidos (CMS) son un tipo de software que nos permiten crear y editar "
    "páginas web sin la necesidad de programar, solo usando la intuición para crear un producto final."
)
add_paragraph(p61)

# --- 6.2 ---
add_heading('6.2 Diagnóstico del CMS actual de Farmacia Carol', level=2)

p62 = (
    "Farmacia Carol no usa un CMS moderno. Su sitio web está construido en un sistema viejo y obsoleto: "
    "utilizan SharePoint 2013 (descontinuado por Microsoft) para las páginas informativas y VevoCart "
    "(plataforma de 2014 prácticamente muerta) para la tienda online. Por otro lado, un problema mayor es "
    "que no cuenta con integraciones modernas y no es responsive, a pesar de que el 70% de los usuarios en "
    "República Dominicana prefiere utilizar el formato móvil. En el copyright del footer dice 2017 (señal "
    "clara de abandono) y cualquier cambio mínimo requiere un programador que sepa SharePoint. Farmacia Carol "
    "no solo tiene un problema de diseño, sino más bien de infraestructura. La solución más factible es "
    "migrar a un CMS moderno."
)
add_paragraph(p62)

# --- 6.3 ---
add_heading('6.3 Evaluación de opciones de CMS', level=2)

p63_intro = (
    "Para reemplazar el SharePoint + VevoCart se evaluaron las 4 opciones más viables del mercado."
)
add_paragraph(p63_intro)

# Tabla comparativa CMS
table = doc.add_table(rows=7, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

headers = ['Criterio', 'WordPress + WooCommerce', 'Shopify Plus', 'Magento', 'Webflow + Foxy']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

rows_data = [
    ['Facilidad de uso', 'Media', 'Alta', 'Baja', 'Media'],
    ['Escalabilidad (+5,000 productos)', 'Alta', 'Muy alta', 'Muy alta', 'Media'],
    ['Costo inicial', 'Bajo ($500-2,000)', 'Medio ($2,000/mes)', 'Muy alto ($22,000+/año)', 'Bajo ($1,000)'],
    ['E-commerce nativo', 'Sí (con WooCommerce)', 'Sí (líder mundial)', 'Sí (enterprise)', 'Limitado'],
    ['SEO técnico', 'Excelente', 'Bueno', 'Excelente', 'Excelente'],
    ['Comunidad y soporte en RD', 'Muy amplia', 'Amplia', 'Limitada', 'Escasa'],
]

for i, row_data in enumerate(rows_data, start=1):
    for j, val in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

doc.add_paragraph()

add_paragraph('Prompt utilizado para generar la tabla comparativa (herramienta: Claude):', bold=True)

prompt_tabla = (
    "\"Genera una tabla comparativa de 4 CMS para un proyecto de migración de una farmacia dominicana con "
    "4,856 productos y 60+ sucursales (actualmente en SharePoint 2013 + VevoCart). Las opciones a comparar "
    "son: WordPress + WooCommerce, Shopify Plus, Magento (Adobe Commerce) y Webflow + Foxy. Evalúa cada una "
    "en 6 criterios: facilidad de uso, escalabilidad para +5,000 productos, costo inicial, e-commerce nativo, "
    "SEO técnico y comunidad/soporte en República Dominicana. Formato markdown.\""
)
p = doc.add_paragraph()
run = p.add_run(prompt_tabla)
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

p63_final = (
    "La IA generó la estructura y valores iniciales de la tabla. Luego se revisaron y ajustaron manualmente "
    "los datos de costos y la columna de \"comunidad en RD\" según el conocimiento real del mercado local."
)
add_paragraph(p63_final)

# --- 6.4 ---
add_heading('6.4 Recomendación final: WordPress + WooCommerce', level=2)

p64_1 = (
    "Para Farmacia Carol la mejor opción es WordPress + WooCommerce. Es la que más se ajusta a lo que la "
    "farmacia necesita hoy: mueve sin problema los más de 4,800 productos del catálogo, maneja el blog, las "
    "páginas de sucursales, los servicios y el programa Caroleal todo en un mismo sitio, y se puede conectar "
    "fácil con WhatsApp, Azul, CardNet y todas las herramientas que se usan en República Dominicana."
)
add_paragraph(p64_1)

p64_2 = (
    "Además, es mucho más barato que Shopify Plus (se ahorran como $20,000 el primer año) y en RD hay un "
    "montón de gente que trabaja con WordPress, así que Farmacia Carol no depende de una sola persona ni de "
    "un solo proveedor. Al final, es una plataforma moderna, fácil de actualizar y con espacio para crecer "
    "sin tener que cambiar de sistema otra vez dentro de 10 años."
)
add_paragraph(p64_2)

# --- 6.5 ---
add_heading('6.5 Arquitectura propuesta del sitio (Mockups HD)', level=2)

p65_intro = (
    "Para mostrar cómo se vería Farmacia Carol con la nueva plataforma (WordPress + WooCommerce), se "
    "crearon 5 mockups de alta fidelidad de las pantallas clave. Estos son el paso siguiente a los wireframes "
    "del punto 5.4: ahora con colores institucionales (rojo Carol), tipografía real, fotos de productos y "
    "el diseño exacto que se vería al visitar el sitio."
)
add_paragraph(p65_intro)

p65_note = (
    "(Mockups generados con IA — herramienta: nano-banana-images / Kie.ai con Gemini nano-banana-2. Ver prompts completos en la Declaración de Uso de IA al final del documento.)"
)
p = doc.add_paragraph()
run = p.add_run(p65_note)
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

hd_files = [
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/hd-mockups/hd1-home.jpg', 'Figura 7. Mockup HD — Home.'),
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/hd-mockups/hd2-product.jpg', 'Figura 8. Mockup HD — Ficha de producto.'),
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/hd-mockups/hd3-category.jpg', 'Figura 9. Mockup HD — Categoría con filtros.'),
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/hd-mockups/hd4-sucursales.jpg', 'Figura 10. Mockup HD — Sucursales.'),
    ('/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/hd-mockups/hd5-checkout.jpg', 'Figura 11. Mockup HD — Checkout.'),
]

for path, caption in hd_files:
    if add_image_if_exists(path, width_inches=5):
        add_caption(caption)
    else:
        add_paragraph(f'[INSERTAR IMAGEN: {os.path.basename(path)}]', bold=True)
        add_caption(caption)

# --- 6.6 ---
add_heading('6.6 Plan de implementación y presupuesto', level=2)

p66_1 = (
    "Migrar de SharePoint 2013 + VevoCart a WordPress + WooCommerce es un proyecto que requiere planificación "
    "por fases para no afectar las ventas actuales ni perder datos del catálogo. A continuación se presenta "
    "el plan de implementación propuesto para Farmacia Carol."
)
add_paragraph(p66_1)

fases = [
    "Descubrimiento y planificación (2 semanas).",
    "Diseño y prototipado del wireframes y mockups aprobados por el cliente (3 semanas).",
    "Desarrollo en WordPress + WooCommerce y setup del hosting (6 semanas).",
    "Migración de contenido (2 semanas).",
    "Pruebas y control de calidad (2 semanas).",
    "Lanzamiento (1 semana).",
    "Mantenimiento (continuo).",
]

for f in fases:
    p = doc.add_paragraph(f, style='List Number')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

add_paragraph('Tiempo total estimado: 16 semanas (4 meses) desde el descubrimiento hasta el lanzamiento.', bold=True)

doc.add_paragraph()
add_paragraph('Presupuesto de la inversión inicial:', bold=True)

# Tabla de inversión inicial
table2 = doc.add_table(rows=14, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_ALIGN_PARAGRAPH.CENTER

inv_data = [
    ['Concepto', 'Costo (USD)'],
    ['Descubrimiento y planificación', '$800'],
    ['Diseño UX/UI (mockups, sistema de diseño)', '$1,500'],
    ['Desarrollo WordPress + tema personalizado', '$3,500'],
    ['Integración e-commerce (WooCommerce, pagos, envíos)', '$1,200'],
    ['Migración de datos (productos, sucursales, contenido)', '$800'],
    ['Testing y QA', '$500'],
    ['Lanzamiento y configuración DNS/SEO', '$400'],
    ['Subtotal desarrollo', '$8,700'],
    ['Licencias de plugins premium (Elementor Pro, RankMath Pro, WP Rocket)', '$500'],
    ['Hosting premium año 1 (Cloudways o Kinsta)', '$600'],
    ['Dominio y certificado SSL', '$50'],
    ['Subtotal primer año', '$1,150'],
    ['TOTAL INVERSIÓN INICIAL', '$9,850'],
]

for i, row in enumerate(inv_data):
    for j, val in enumerate(row):
        cell = table2.rows[i].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0 or 'Subtotal' in val or 'TOTAL' in val:
                    run.bold = True

doc.add_paragraph()
add_paragraph('Costos recurrentes desde el año 2:', bold=True)

# Tabla costos recurrentes
table3 = doc.add_table(rows=5, cols=2)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_ALIGN_PARAGRAPH.CENTER

rec_data = [
    ['Concepto', 'Costo anual (USD)'],
    ['Hosting premium', '$600'],
    ['Licencias de plugins', '$500'],
    ['Mantenimiento técnico (actualizaciones, backups, seguridad)', '$1,200'],
    ['TOTAL anual', '$2,300'],
]

for i, row in enumerate(rec_data):
    for j, val in enumerate(row):
        cell = table3.rows[i].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0 or 'TOTAL' in val:
                    run.bold = True

doc.add_page_break()

# ============================================================
# CONCLUSIÓN
# ============================================================
add_heading('Conclusión', level=1)

conclusion = (
    "Farmacia Carol es una marca dominante en el mercado farmacéutico dominicano con más de 48 años de "
    "trayectoria y más de 60 sucursales activas. Sin embargo, su infraestructura web no está a la altura "
    "de su reputación. A lo largo del Módulo 5 se diagnosticó que el sitio actual combina dos sistemas "
    "obsoletos (SharePoint 2013 y VevoCart) repartidos en cinco subdominios inconexos, con un catálogo de "
    "4,856 productos sin filtros, páginas de producto pobres, contenido desactualizado desde el 2017, sin "
    "integración de WhatsApp y una experiencia móvil deficiente."
)
add_paragraph(conclusion)

c2 = (
    "La propuesta desarrollada en este trabajo ofrece una solución integral: una nueva arquitectura "
    "jerárquica de 3 niveles bajo un solo dominio, wireframes y mockups HD que validan visualmente la "
    "estructura, y la recomendación de migrar a WordPress + WooCommerce como plataforma CMS. Esta opción "
    "ofrece el mejor balance entre costo, flexibilidad, escalabilidad y soporte local en República "
    "Dominicana, con una inversión inicial de $9,850 USD y un mantenimiento anual de $2,300 USD."
)
add_paragraph(c2)

c3 = (
    "Los beneficios esperados son claros: un sitio web unificado, moderno y rápido; más de 5,000 páginas "
    "indexables en Google; aparición con rich data en Google Maps para las 60+ sucursales; integración "
    "nativa con WhatsApp Business y las pasarelas de pago locales (Azul, CardNet); y autonomía total del "
    "equipo interno para gestionar el contenido sin depender de un programador. En definitiva, Farmacia "
    "Carol tiene la oportunidad de convertir su presencia web en un verdadero canal de ventas a la altura "
    "de su liderazgo en el mercado físico."
)
add_paragraph(c3)

doc.add_page_break()

# ============================================================
# BIBLIOGRAFÍA
# ============================================================
add_heading('Bibliografía', level=1)

biblio = [
    "Farmacia Carol. (2024). Sitio web oficial. https://www.farmaciacarol.com/",
    "W3Techs. (2025). Usage statistics of content management systems. https://w3techs.com/technologies/overview/content_management",
    "WordPress.org. (2025). Documentación oficial. https://wordpress.org/documentation/",
    "WooCommerce. (2025). WooCommerce documentation. https://woocommerce.com/documentation/",
    "Google Developers. (2025). Search Central — Documentación de SEO. https://developers.google.com/search/docs",
    "Nielsen Norman Group. (2020). The 3-Click Rule for Navigation Is False. https://www.nngroup.com/articles/3-click-rule/",
    "Krug, S. (2014). Don't Make Me Think, Revisited: A Common Sense Approach to Web Usability. New Riders.",
    "Shopify. (2025). Shopify Plus vs WooCommerce. https://www.shopify.com/plus",
    "Cloudways. (2025). Managed WordPress Hosting. https://www.cloudways.com/",
    "Adobe Commerce (Magento). (2025). Adobe Commerce pricing and features. https://business.adobe.com/products/magento/magento-commerce.html",
]

for b in biblio:
    p = doc.add_paragraph(b)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.left_indent = Cm(1.25)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# DECLARACIÓN DE USO DE IA
# ============================================================
add_heading('Declaración de Uso de Inteligencia Artificial', level=1)

decl = (
    "En cumplimiento con los lineamientos del curso, se declara el uso de herramientas de Inteligencia "
    "Artificial en secciones específicas de este trabajo. A continuación se detallan las herramientas "
    "utilizadas, las secciones donde se aplicaron y los prompts exactos utilizados."
)
add_paragraph(decl)

add_heading('Herramientas utilizadas', level=2)

tools = [
    "Claude (Anthropic) — asistencia en redacción de secciones puntuales y generación de la tabla comparativa de CMS.",
    "Excalidraw-diagram (skill de Claude) — generación del diagrama del mapa de sitio propuesto (punto 5.3).",
    "nano-banana-images vía Kie.ai (modelo Gemini nano-banana-2) — generación de los 5 wireframes de baja fidelidad (punto 5.4) y los 5 mockups de alta fidelidad (punto 6.5).",
]

for t in tools:
    p = doc.add_paragraph(t, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

add_heading('Prompt del diagrama del mapa de sitio (Punto 5.3)', level=2)
p = doc.add_paragraph()
run = p.add_run(prompt_excalidraw)
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

add_heading('Prompts de los wireframes (Punto 5.4)', level=2)
add_paragraph('Herramienta: nano-banana-images (Kie.ai con Gemini). Configuración común: aspect_ratio 3:4, resolution 1K, output_format jpg.', size=10)

wireframe_prompts = [
    ("Wireframe 1 — Home (wf1-home.jpg)",
     "Low fidelity website wireframe mockup, black and white line art, clean minimal grid layout, desktop resolution. Website home page wireframe for a pharmacy chain named FARMACIA CAROL. Vertical layout on white background with thin black borders and placeholder gray boxes. TOP: Header bar with 'FARMACIA CAROL' logo text on left, navigation menu items: Productos, Sucursales, Servicios, Caroleal, Blog, Contacto, search bar, and 3 icons on right (user, cart, whatsapp). BELOW HEADER: Large hero banner rectangle with X marks inside and text 'HERO BANNER - Promociones rotativas'. NEXT SECTION: Horizontal row of 6 category cards each with icon placeholder and labels: Medicamentos, Salud, Dermocosmetica, Cuidado Personal, Bebe, Minimarket. NEXT SECTION: Title 'MAS VENDIDOS' with grid of 8 product cards, each showing image placeholder, product name lines, price, and Add to cart button. NEXT SECTION: 3 large service blocks side by side labeled 'BUSCAR DOCTOR', 'CONSULTA AL FARMACEUTICO', 'AYUDAMED' each with icon and CTA button. NEXT SECTION: Store locator with map placeholder on left and 3 branch cards on right. BOTTOM: Footer with 3 columns showing logo, links and contact info. Annotations with arrows and handwritten notes pointing to elements. Pure wireframe style, no colors except black, white and gray fills, technical blueprint look, clean sans serif font like Helvetica, minimal design documentation style."),

    ("Wireframe 2 — Producto (wf2-product.jpg)",
     "Low fidelity website wireframe mockup, black and white line art, clean minimal grid layout, desktop resolution. Product detail page wireframe for FARMACIA CAROL pharmacy website. White background with thin black borders and placeholder gray boxes with X marks. TOP: Header bar with FARMACIA CAROL logo, navigation menu, search bar, user icon, cart icon. BELOW HEADER: Breadcrumbs text 'Inicio > Productos > Medicamentos > Analgesicos > Acetaminofen 500mg'. MAIN SECTION split in 2 columns. LEFT COLUMN 60% width: Large product image placeholder square with X, 4 small thumbnail squares below it, red label tag on top that says 'REQUIERE RECETA'. RIGHT COLUMN 40% width: Product title lines, row of 5 stars, crossed out old price and new big price with discount badge '20% OFF', quantity selector with minus plus buttons, large dark button 'AGREGAR AL CARRITO', secondary button 'COMPRAR AHORA', box labeled 'Disponibilidad por sucursal' with list of 3 branches, share buttons for WhatsApp, Facebook, Email. BELOW THE 2 COLUMNS: Tabs row with 4 tabs: Descripcion, Modo de Uso, Ingredientes, FAQ. NEXT SECTION: 'Productos Relacionados' title with 4 product cards. BOTTOM: Reviews section with stars and comment boxes. FOOTER at the end. Annotations with arrows pointing to key elements, handwritten style notes. Pure wireframe blueprint style, only black white and gray, clean sans serif typography, technical documentation look."),

    ("Wireframe 3 — Categoría (wf3-category.jpg)",
     "Low fidelity website wireframe mockup, black and white line art, clean minimal grid layout, desktop resolution. Product category listing page wireframe for FARMACIA CAROL pharmacy website. White background with thin black borders and gray placeholder boxes with X marks. TOP: Header bar with FARMACIA CAROL logo, navigation menu, search bar, icons. Below header: Breadcrumbs 'Inicio > Productos > Medicamentos'. Banner strip with category title 'MEDICAMENTOS' and short description line. MAIN SECTION split in 2 columns. LEFT COLUMN 25% width sidebar labeled 'FILTROS' with sections: Subcategorias with checkboxes and counters, Marcas with checkboxes and numbers, Precio with slider bar, Forma farmaceutica with checkboxes (Tabletas, Jarabe, Capsulas, Crema), Disponibilidad checkboxes, Con o sin receta checkboxes. RIGHT COLUMN 75% width: Top bar with text '4856 productos encontrados', dropdown 'Ordenar por: Relevancia', toggle icons for grid and list view. BELOW: Grid of 12 product cards in 4 columns. Each card shows image placeholder square with X, product name 2 lines, price, small Add to cart button. BOTTOM: Pagination with page numbers 1 2 3 4 5. Footer at the very bottom. Annotations with arrows pointing to filter section and product cards, handwritten notes. Pure wireframe blueprint style, only black white gray, clean sans serif typography, technical UX documentation look."),

    ("Wireframe 4 — Sucursales (wf4-sucursales.jpg)",
     "Low fidelity website wireframe mockup, black and white line art, clean minimal grid layout, desktop resolution. Store locator page wireframe for FARMACIA CAROL pharmacy website. White background with thin black borders and gray placeholder boxes. TOP: Header bar with FARMACIA CAROL logo, navigation menu, search bar, icons. Breadcrumbs 'Inicio > Sucursales'. Large title 'NUESTRAS SUCURSALES' with subtitle 'Encuentra la farmacia mas cerca de ti'. BELOW TITLE: Full width search bar with 3 elements: text input 'Buscar por ciudad', dropdown 'Provincia', dropdown 'Servicio (24 horas, Delivery, Pick-up)', and a button on the right 'Usar mi ubicacion' with location pin icon. MAIN SECTION split in 2 columns. LEFT COLUMN 35% width labeled 'LISTADO DE SUCURSALES': 5 branch cards stacked vertically. Each card contains: branch name in bold, address line, distance text '2.3 km', current status badge 'Abierta ahora' in green or 'Cerrada' in red, horario text, row of small service icons (24h, delivery, farmaceutico, pickup), and 2 buttons at bottom 'Llamar' and 'WhatsApp'. RIGHT COLUMN 65% width: Large full interactive map placeholder of Dominican Republic with multiple pin markers scattered across it, one pin with popup bubble showing branch info. At the top right corner of map: zoom in and zoom out buttons. Footer at bottom. Annotations with arrows pointing to search filters and map pins, handwritten style notes. Pure wireframe blueprint style, only black white gray, clean sans serif typography."),

    ("Wireframe 5 — Checkout (wf5-checkout.jpg)",
     "Low fidelity website wireframe mockup, black and white line art, clean minimal grid layout, desktop resolution. Checkout page wireframe for FARMACIA CAROL pharmacy website. White background with thin black borders and gray placeholder boxes. TOP: Header bar with FARMACIA CAROL logo only, minimal navigation. BELOW HEADER: Progress bar with 3 steps shown as circles connected by lines: 1 CARRITO checked, 2 DATOS DE ENTREGA highlighted with dark circle, 3 PAGO empty circle. Step 2 label shown in bold. MAIN SECTION split in 2 columns with large gap. LEFT COLUMN 65% width labeled 'DATOS DE ENTREGA': Radio buttons at top 'Delivery a domicilio' selected and 'Retiro en sucursal'. Form fields stacked: Nombre completo input, Telefono input, Direccion input long, Sector input, Ciudad dropdown, Referencias textarea. Below form: Textarea labeled 'Notas especiales al pedido'. Checkbox 'Acepto terminos y condiciones'. RIGHT COLUMN 35% width is a sticky box labeled 'RESUMEN DEL PEDIDO' with border: 3 product rows each with small image thumbnail square, name, quantity, price. Divider line. Rows for Subtotal, Delivery, Descuento Caroleal, TOTAL in bold and bigger. Large dark button at bottom 'CONTINUAR AL PAGO'. Footer minimal at bottom. Annotations with arrows pointing to progress bar and sticky order summary, handwritten style notes. Pure wireframe blueprint style, only black white gray, clean sans serif typography, technical UX documentation look."),
]

for title, prompt in wireframe_prompts:
    add_paragraph(title, bold=True, size=11)
    p = doc.add_paragraph()
    run = p.add_run(prompt)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)

add_heading('Prompts de los mockups HD (Punto 6.5)', level=2)
add_paragraph('Herramienta: nano-banana-images (Kie.ai con Gemini nano-banana-2). Configuración común: aspect_ratio 3:4, resolution 2K, output_format jpg.', size=10)

hd_prompts = [
    ("Mockup HD 1 — Home (hd1-home.jpg)",
     "Ultra high fidelity modern e-commerce website mockup, Figma-style UI design, desktop resolution 1920x1080, pixel perfect, clean professional layout. Home page for FARMACIA CAROL (Dominican Republic pharmacy chain). Brand colors: primary red #E30613, white background, dark charcoal text #1A1A1A, accent green #00A651. TOP HEADER: White sticky header with FARMACIA CAROL logo in red on left (cross symbol + text), horizontal navigation menu: Productos, Sucursales, Servicios, Caroleal, Blog, Contacto. On right: search bar with magnifying glass icon, user profile icon, shopping cart icon with badge showing '2', green WhatsApp button. BELOW HEADER: Large hero banner with red gradient background, left side shows bold white headline 'TU SALUD, NUESTRA PRIORIDAD', subheadline 'Envios a domicilio en 2 horas', white CTA button 'COMPRAR AHORA'. Right side of hero: photorealistic image of smiling Dominican pharmacist in white coat holding tablet. BELOW HERO: Horizontal row of 6 category cards with rounded corners, soft shadows, each with a colored icon (pill bottle for Medicamentos, heart for Salud, cream tube for Dermocosmetica, soap for Cuidado Personal, baby bottle for Bebe, shopping basket for Minimarket) and label text below. NEXT SECTION: 'MAS VENDIDOS' title in bold black, grid of 8 product cards with real product photos on white background, each card shows product image, brand name in gray, product name in black, old price crossed out, new price in red bold, red 'Agregar' button. NEXT SECTION: 3 large service blocks with illustrated icons side by side: 'BUSCAR DOCTOR' with stethoscope icon, 'CONSULTA AL FARMACEUTICO' with chat bubble, 'AYUDAMED' with pill reminder icon, each with short description and red outline button. NEXT SECTION: Store locator with interactive Google Maps style map showing Dominican Republic with red pharmacy pins, on the right 3 branch cards showing branch name, address, open status in green, and call/whatsapp buttons. FOOTER: Dark red background with 4 columns: logo and social icons, Productos links, Servicios links, Contacto info with phone and email. Modern sans serif typography (Inter or Poppins font). Soft drop shadows, rounded corners, professional Figma mockup aesthetic, clean spacing, photorealistic, high definition, 4K quality."),

    ("Mockup HD 2 — Producto (hd2-product.jpg)",
     "Ultra high fidelity modern e-commerce product detail page mockup, Figma-style UI design, desktop resolution 1920x1080, pixel perfect, clean professional layout. Product detail page for FARMACIA CAROL pharmacy. Brand colors: primary red #E30613, white background, dark charcoal text #1A1A1A, accent green #00A651. TOP HEADER: White sticky header with FARMACIA CAROL red logo, navigation menu Productos Sucursales Servicios Caroleal Blog Contacto, search bar, user icon, cart icon with badge, green WhatsApp button. BELOW HEADER: Breadcrumbs in light gray 'Inicio > Productos > Medicamentos > Analgesicos > Acetaminofen 500mg'. MAIN SECTION split in 2 columns with 60/40 ratio. LEFT COLUMN: Large product photo of Acetaminofen 500mg medicine box, photorealistic pharmaceutical product photography on white background with soft shadow, small red badge label in top left 'REQUIERE RECETA', below main image row of 4 small thumbnail images showing different angles. RIGHT COLUMN: Small gray text 'GENOMMA LAB', product title in large bold black 'Acetaminofen 500mg - 20 Tabletas', row of 5 golden stars and text '(127 resenas)', crossed out old price 'RD$129.00' in gray, new big price 'RD$99.00' in red bold, small red badge '20% OFF', text 'IVA incluido' in gray. Below: label 'Cantidad' with minus/plus quantity selector showing 1. Large red button 'AGREGAR AL CARRITO' full width with cart icon. Secondary outline button 'COMPRAR AHORA'. Box with light gray border labeled 'Disponibilidad por sucursal': list of 3 branches each with green dot and 'En stock' status and distance. Row of share buttons: green WhatsApp, blue Facebook, gray email. BELOW THE 2 COLUMNS: Tabs row with 4 tabs: Descripcion (active with red underline), Modo de Uso, Ingredientes, Preguntas Frecuentes. Tab content showing product description paragraph. NEXT SECTION: 'Productos Relacionados' title bold, horizontal carousel of 4 product cards with photos, names and prices in red. BOTTOM: Reviews section with average rating, progress bars for 5-4-3-2-1 stars, 3 review cards with user names and comments. FOOTER: Dark red with columns. Modern sans serif typography Inter font, soft shadows, rounded corners, professional Figma mockup, photorealistic, 4K HD quality."),

    ("Mockup HD 3 — Categoría (hd3-category.jpg)",
     "Ultra high fidelity modern e-commerce category listing page mockup, Figma-style UI design, desktop resolution 1920x1080, pixel perfect, clean professional layout. Category page MEDICAMENTOS for FARMACIA CAROL pharmacy. Brand colors: primary red #E30613, white background, dark charcoal text #1A1A1A, accent green #00A651. TOP HEADER: White sticky header with FARMACIA CAROL red logo, navigation menu, search bar, icons, green WhatsApp button. BELOW HEADER: Breadcrumbs 'Inicio > Productos > Medicamentos'. Red banner strip with white bold title 'MEDICAMENTOS' and subtitle 'Encuentra todos los medicamentos que necesitas'. MAIN SECTION split in 2 columns. LEFT COLUMN 25% width sidebar with white background and subtle border labeled 'FILTROS' in bold red: Section 'Subcategorias' with checkboxes: Dolor y Fiebre (125), Alergias (87), Antibioticos (60), Digestivo (54), Vitaminas (198). Section 'Marcas' with checkboxes: Genomma, Bayer, Pfizer, GSK, Roche. Section 'Precio' with range slider showing RD$50 - RD$500. Section 'Forma farmaceutica' with checkboxes: Tabletas, Jarabe, Capsulas, Crema, Inyectable. Section 'Disponibilidad' with checkbox 'En stock'. Section 'Con o sin receta' with radio buttons. Red 'APLICAR FILTROS' button at bottom. RIGHT COLUMN 75% width: Top bar with text '4,856 productos encontrados' on left, dropdown 'Ordenar por: Relevancia' on right, grid/list view toggle icons. Grid of 12 product cards in 4 columns. Each card has white background with soft shadow and rounded corners, showing photorealistic product image of real medicine packages, small gray brand name at top, product name 2 lines, row of small stars, price in red bold, red 'Agregar' button with cart icon. Products shown: Acetaminofen, Ibuprofeno, Aspirina, Loratadina, Amoxicilina, Omeprazol, Paracetamol, Diclofenaco, Naproxeno, Metformina, Losartan, Atorvastatina. BOTTOM: Pagination with numbered buttons 1 2 3 4 5 with page 1 highlighted in red. FOOTER: Dark red background with columns. Modern sans serif Inter typography, soft shadows, rounded corners, professional Figma mockup, photorealistic pharmaceutical products, 4K HD quality."),

    ("Mockup HD 4 — Sucursales (hd4-sucursales.jpg)",
     "Ultra high fidelity modern store locator page mockup, Figma-style UI design, desktop resolution 1920x1080, pixel perfect, clean professional layout. Store locator page for FARMACIA CAROL pharmacy chain in Dominican Republic. Brand colors: primary red #E30613, white background, dark charcoal text #1A1A1A, accent green #00A651. TOP HEADER: White sticky header with FARMACIA CAROL red logo, navigation menu Productos Sucursales Servicios Caroleal Blog Contacto, search bar, icons, green WhatsApp button. BELOW HEADER: Breadcrumbs 'Inicio > Sucursales'. Large centered title 'NUESTRAS SUCURSALES' in bold black with red underline accent, subtitle 'Encuentra la farmacia mas cerca de ti' in gray. SEARCH BAR SECTION: Full width card with white background and soft shadow containing 4 elements in a row: text input with search icon and placeholder 'Buscar por ciudad o sector', dropdown 'Provincia' showing 'Santo Domingo', dropdown 'Servicio' showing '24 horas, Delivery, Pick-up', red button 'USAR MI UBICACION' with location pin icon. MAIN SECTION split in 2 columns. LEFT COLUMN 35% width labeled 'LISTADO DE SUCURSALES (62)': scrollable list of 5 branch cards stacked vertically. Each card has white background soft shadow rounded corners: branch name in bold black 'Farmacia Carol Naco', address line in gray 'Av. Tiradentes 52, Santo Domingo', distance text '2.3 km' with pin icon, current status badge 'Abierta ahora' in green pill shape, horario text '8:00 AM - 10:00 PM', row of small colored service icons (24h in red, delivery truck in blue, pharmacist in green, pickup in orange), 2 buttons at bottom 'Llamar' outline red with phone icon and 'WhatsApp' filled green. RIGHT COLUMN 65% width: Large interactive Google Maps style map placeholder showing Dominican Republic island with multiple red pharmacy pin markers scattered across Santo Domingo, Santiago, La Vega, and other cities. One pin with popup bubble showing branch name, address, status and mini photo. Top right corner of map: floating zoom in/out buttons and location centering button. FOOTER: Dark red background with columns. Modern sans serif Inter typography, soft shadows, rounded corners, professional Figma mockup, photorealistic map, 4K HD quality."),

    ("Mockup HD 5 — Checkout (hd5-checkout.jpg)",
     "Ultra high fidelity modern e-commerce checkout page mockup, Figma-style UI design, desktop resolution 1920x1080, pixel perfect, clean professional layout. Checkout page for FARMACIA CAROL pharmacy website. Brand colors: primary red #E30613, white background, dark charcoal text #1A1A1A, accent green #00A651. TOP HEADER: White minimal header with FARMACIA CAROL red logo centered, small lock icon and text 'CHECKOUT SEGURO' on right. BELOW HEADER: Progress bar with 3 steps shown as circles connected by lines with labels below: Step 1 'CARRITO' with green checkmark circle, Step 2 'DATOS DE ENTREGA' with red filled circle highlighted and bold label, Step 3 'PAGO' with empty gray circle. MAIN SECTION split in 2 columns with large gap. LEFT COLUMN 60% width labeled 'DATOS DE ENTREGA' in bold: 2 radio button cards at top side by side, 'Delivery a domicilio' selected with red border and truck icon, 'Retiro en sucursal' unselected with store icon. Form fields stacked below with floating labels and clean borders: Nombre completo input with user icon, Telefono input with phone icon showing '+1 809 555 1234', Direccion input long with map pin icon, Sector input, Ciudad dropdown showing 'Santo Domingo', Referencias textarea 'Ej: Casa amarilla con porton negro'. Separator line. Textarea labeled 'Notas especiales al pedido' with placeholder 'Alguna indicacion especial?'. Checkbox with 'Acepto terminos y condiciones y politica de privacidad'. RIGHT COLUMN 40% width is a sticky card with white background soft shadow and subtle border labeled 'RESUMEN DEL PEDIDO' in bold: 3 product rows each with small photorealistic product thumbnail (Acetaminofen, Vitamina C, Protector Solar), product name in black, quantity in gray 'Cant: 2', price in black 'RD$198.00'. Divider line. Summary rows: Subtotal RD$547.00, Delivery RD$150.00, Descuento Caroleal -RD$50.00 in green, separator line, TOTAL RD$647.00 in bold bigger red. Large red button full width 'CONTINUAR AL PAGO' with lock icon. Below button: row of accepted payment logos (Visa, Mastercard, Azul, CardNet). Small security badges row: SSL seguro, Pago protegido. FOOTER: Minimal dark footer with contact info. Modern sans serif Inter typography, soft shadows, rounded corners, professional Figma mockup, photorealistic products, 4K HD quality."),
]

for title, prompt in hd_prompts:
    add_paragraph(title, bold=True, size=11)
    p = doc.add_paragraph()
    run = p.add_run(prompt)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)

add_heading('Revisión humana', level=2)
rev = (
    "Todos los textos redactados fueron revisados, editados y validados por el estudiante. Las imágenes "
    "generadas por IA (diagrama, wireframes y mockups HD) fueron revisadas visualmente y aceptadas como "
    "representaciones ilustrativas de las propuestas. Los datos de mercado, costos, plazos y decisiones "
    "estratégicas fueron contrastados con el análisis personal del autor como consultor de mercadeo digital "
    "en el contexto de la República Dominicana."
)
add_paragraph(rev)

# ============================================================
# GUARDAR
# ============================================================
output_path = '/Users/martinmercedes/Desktop/Executive assistant 2/projects/universidad/Trabajo-Final-SIG-Farmacia-Carol.docx'
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
