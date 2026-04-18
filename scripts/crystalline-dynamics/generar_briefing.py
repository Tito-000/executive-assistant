#!/usr/bin/env python3
"""
Genera el briefing completo de Crystalline Dynamics (Andri Ramírez).
Un solo documento Word con toda la información que necesitamos de ella.
MM Agency se encarga de crear GMB, Google Ads, Analytics, dominio, etc.

Uso: python3 generar_briefing.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from pathlib import Path

W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

OUT_DIR = Path(__file__).resolve().parents[2] / "outputs" / "briefings" / "crystalline-dynamics"
OUT_DIR.mkdir(parents=True, exist_ok=True)

YELLOW = RGBColor(0x9A, 0x7A, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
LINE_GREY = RGBColor(0xBB, 0xBB, 0xBB)


def set_margins(doc, top=0.8, bottom=0.8, left=0.9, right=0.9):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_header_block(doc, subtitle):
    p = doc.add_paragraph()
    run = p.add_run("MM AGENCY")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = YELLOW
    p.paragraph_format.space_after = Pt(0)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Crystalline Dynamics · Fence Growth System · Abril 2026")
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY
    p2.paragraph_format.space_after = Pt(18)

    pt = doc.add_paragraph()
    rt = pt.add_run(subtitle)
    rt.bold = True
    rt.font.size = Pt(20)
    rt.font.color.rgb = BLACK
    pt.paragraph_format.space_after = Pt(6)


def add_intro(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(18)


def add_section(doc, num, title):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{num}.  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = YELLOW
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def _build_sdt(placeholder):
    xml = f'''<w:sdt {W_NS}>
      <w:sdtPr>
        <w:rPr><w:rStyle w:val="PlaceholderText"/></w:rPr>
        <w:alias w:val="respuesta"/>
        <w:tag w:val="respuesta"/>
        <w:showingPlcHdr/>
        <w:text w:multiLine="1"/>
      </w:sdtPr>
      <w:sdtContent>
        <w:r>
          <w:rPr><w:rStyle w:val="PlaceholderText"/><w:color w:val="888888"/></w:rPr>
          <w:t>{placeholder}</w:t>
        </w:r>
      </w:sdtContent>
    </w:sdt>'''
    return parse_xml(xml)


def _add_border_box(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:color'), 'BBBBBB')
        b.set(qn('w:space'), '6')
        pBdr.append(b)
    pPr.append(pBdr)


def add_question(doc, text, lines=2, placeholder="✎  Escribe tu respuesta aquí (selecciona este texto y sobrescribe)"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(4)

    box = doc.add_paragraph()
    _add_border_box(box)
    box.paragraph_format.space_after = Pt(12)
    rb = box.add_run(placeholder)
    rb.italic = True
    rb.font.size = Pt(10)
    rb.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"⚑  {text}")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(2)


def add_footer_contact(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("— — — — — — — — — — — — — — — — — — —")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Martin Mercedes · MM Agency\n")
    r2.bold = True
    r2.font.size = Pt(10)
    r3 = p2.add_run("martin@mmagency.do · WhatsApp directo")
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = GREY
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ========================================================================
# BRIEFING COMPLETO — TODO EN UN SOLO DOC
# ========================================================================
def build_briefing():
    doc = Document()
    set_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    add_header_block(doc, "Briefing del Proyecto")

    add_intro(
        doc,
        "Hola Andri — para arrancar con Crystalline Fence necesito que me respondas estas "
        "preguntas y me compartas los materiales que te pido al final. No tienes que escribir "
        "mucho; respuestas cortas y directas están bien. Si alguna no aplica o aún no la tienes "
        "decidida, déjala en blanco y lo vemos juntos en el kickoff call. "
        "Nosotros nos encargamos de crear el Google Business Profile, Google Ads, Analytics, "
        "dominio y todo lo técnico — de tu lado solo necesitamos la info del negocio "
        "y los materiales (fotos, licencias, reseñas).",
    )

    # ========== PARTE 1: NOMBRE & IDENTIDAD ==========
    add_section(doc, 1, "Nombre Comercial y Dominio")
    add_question(
        doc,
        "¿Qué nombre comercial quieres usar para la línea de Fencing? "
        "(ej: \"Crystalline Fence\", \"Crystalline Fence Co.\", u otro). "
        "Si aún no lo tienes decidido, déjalo en blanco y te mando 3 opciones:",
        lines=3,
    )
    add_question(
        doc,
        "¿Ya tienes un dominio en mente? (ej: crystallinefence.com). "
        "Si no, déjalo en blanco — nosotros te mandamos 3 dominios disponibles y tú eliges:",
        lines=2,
    )
    add_note(doc, "Nosotros nos encargamos de comprar, configurar y apuntar el dominio al hosting.")

    # ========== PARTE 2: INFO DEL NEGOCIO ==========
    add_section(doc, 2, "Negocio")
    add_question(doc, "¿Cuántos años llevan operando en fencing?")
    add_question(doc, "¿Cuántas personas hay en el equipo (incluyéndote a ti)?")
    add_question(doc, "¿Cuántas instalaciones de fence pueden hacer por semana con el equipo actual?")
    add_question(
        doc,
        "Dirección física exacta de la oficina/warehouse (la que irá en Google Maps y en la web):",
        lines=2,
    )
    add_question(doc, "Horario de atención (días y horas):", lines=2)

    # ========== PARTE 3: SERVICIOS ==========
    add_section(doc, 3, "Servicios de Fencing")
    add_question(
        doc,
        "¿Qué tipos de fence instalan? (vinyl, wood, aluminum, chain link, pool, commercial, otros — marca todas las que apliquen)",
        lines=3,
    )
    add_question(doc, "¿Cuál es el que MÁS venden hoy?")
    add_question(doc, "¿Cuál les deja MÁS margen (aunque vendan menos)?")
    add_question(
        doc,
        "¿Hacen también repair/reparaciones o solo new install?",
    )

    # ========== PARTE 4: ZONAS ==========
    add_section(doc, 4, "Zonas de Servicio")
    add_question(doc, "¿Qué ciudades/condados cubren desde Davenport? (lista todas)", lines=3)
    add_question(doc, "¿Cuál es el radio máximo en millas que están dispuestos a ir?")
    add_question(doc, "¿Hay zonas donde NO quieren ir? (muy lejos, HOAs difíciles, etc.)")

    # ========== PARTE 5: PRECIOS ==========
    add_section(doc, 5, "Precios")
    add_question(
        doc,
        "Ticket promedio por tipo de fence — rango aproximado (mínimo / promedio / máximo):",
        lines=6,
    )
    add_note(doc, "Esto no se publica en la web sin tu permiso; es para calibrar las campañas y el target del cliente ideal.")

    # ========== PARTE 6: CLIENTE IDEAL ==========
    add_section(doc, 6, "Cliente Ideal")
    add_question(
        doc,
        "¿Quién es el cliente ideal? (marca las que apliquen: homeowner residencial / comercial / HOA / pool builders / realtors / otro)",
        lines=2,
    )
    add_question(
        doc,
        "Describe en 1-2 líneas cómo es tu cliente favorito — el que da menos dolores de cabeza y paga mejor:",
        lines=3,
    )

    # ========== PARTE 7: PROCESO DE VENTA ==========
    add_section(doc, 7, "Proceso de Venta")
    add_question(
        doc,
        "¿Cómo cotizan hoy? (visita en sitio / video call / foto + medidas por WhatsApp / mixto)",
    )
    add_question(doc, "Desde que el lead contacta hasta que le mandas el quote — ¿cuánto tardan en promedio?")
    add_question(doc, "¿Quién contesta las llamadas y mensajes de leads?")
    add_question(doc, "Desde que firman hasta que instalan — ¿cuánto tardan en promedio?")
    add_question(doc, "¿Cuánto cobras de deposit para agendar? (% o monto fijo)")

    # ========== PARTE 8: DIFERENCIADORES ==========
    add_section(doc, 8, "Diferenciadores y Competencia")
    add_question(
        doc,
        "¿Qué hace Crystalline Fence MEJOR que los otros fence contractors de Davenport?",
        lines=3,
    )
    add_question(
        doc,
        "Menciona 3 competidores directos que conozcas en Polk County o Central Florida "
        "(nombre + web o IG si los sabes):",
        lines=3,
    )

    # ========== PARTE 9: FINANCING ==========
    add_section(doc, 9, "Financing")
    add_question(doc, "¿Ofrecen financiamiento? Si sí, ¿con qué plataforma? (Synchrony, Hearth, GreenSky, otro)")
    add_note(doc, "Si no ofrecen, puedo ayudarte a conectar con Hearth o GreenSky — sube mucho la conversión.")

    # ========== PARTE 10: OBJECIONES ==========
    add_section(doc, 10, "Objeciones Comunes")
    add_question(
        doc,
        "¿Qué es lo que MÁS te objetan los clientes antes de firmar? "
        "(precio, tiempo de entrega, materiales, permisos, etc.)",
        lines=3,
    )

    # ========== PARTE 11: ESTACIONALIDAD ==========
    add_section(doc, 11, "Estacionalidad")
    add_question(doc, "¿Qué meses son los MÁS fuertes de ventas?")
    add_question(doc, "¿Qué meses son los más lentos?")
    add_question(
        doc,
        "¿Hay deadlines típicos que empujan a comprar rápido? (pool inspection, HOA notice, mudanza, hurricane season, etc.)",
        lines=2,
    )

    # ========== PARTE 12: META ==========
    add_section(doc, 12, "Meta de los Próximos 3 Meses")
    add_question(
        doc,
        "Si todo sale perfecto, ¿cómo se ve Crystalline Fence en 90 días? "
        "(# de instalaciones/mes, revenue, reputación online)",
        lines=3,
    )

    # ========== PARTE 13: IDENTIDAD VISUAL ==========
    add_section(doc, 13, "Identidad Visual (si ya tienes)")
    add_question(doc, "¿Tienen logo ya hecho? Sí / No. Si sí, mándalo en PNG con fondo transparente o PDF vectorial.")
    add_question(doc, "¿Colores de marca que quieran usar? (ej: azul + gris, o nos mandas códigos hex si tienen)")
    add_question(
        doc,
        "3 webs o cuentas de IG de fencing/construcción cuyo estilo visual te guste "
        "(no tienen que ser competidores, solo referencias):",
        lines=3,
    )
    add_note(doc, "Si NO tienen logo ni colores definidos, no hay problema: trabajamos con lo que proyecte confianza y autoridad local.")

    # ========== PARTE 14: MATERIALES A COMPARTIR ==========
    add_section(doc, 14, "Materiales para Compartir")
    add_intro(
        doc,
        "Esto es lo que necesitamos que nos compartas (por Google Drive, Dropbox, o WhatsApp):",
    )
    add_bullet(doc, "Fotos de trabajos terminados — originales, sin comprimir (mientras más, mejor; ideal before/after)")
    add_bullet(doc, "Fotos del equipo trabajando + camión/truck con logo si lo tienen")
    add_bullet(doc, "Videos de instalaciones o testimoniales (si tienes)")
    add_bullet(doc, "Copia/PDF de la contractor license de Florida")
    add_bullet(doc, "Copia/PDF del liability insurance vigente")
    add_bullet(doc, "Screenshots de las mejores reseñas de Google actuales")
    add_bullet(doc, "Testimonios escritos de clientes (WhatsApp, email) que aún no estén publicados")
    add_bullet(doc, "Logo en alta resolución (si ya tienen)")
    add_question(doc, "Pega aquí el link de la carpeta (Google Drive / Dropbox) o confirma que lo mandas por WhatsApp:", lines=2)

    # ========== PARTE 15: CONTACTO & TELÉFONO ==========
    add_section(doc, 15, "Contacto del Negocio")
    add_question(doc, "Email oficial del negocio (para usar como remitente de formularios de leads):")
    add_question(doc, "Teléfono principal (el que aparecerá en Google, la web y recibirá llamadas de campañas):")
    add_question(doc, "Instagram (@usuario) si tienen:")
    add_question(doc, "Facebook Business Page (URL) si tienen:")
    add_note(doc, "Si no tienen IG o FB aún, nosotros los creamos y configuramos como parte del setup.")

    add_footer_contact(doc)

    out = OUT_DIR / "briefing-crystalline-fence.docx"
    doc.save(out)
    print(f"✓ {out}")
    return out


if __name__ == "__main__":
    build_briefing()
    print("\nListo para enviar a Andri.")
