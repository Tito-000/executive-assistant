"""
Formato Word para que Memorama complete la tabla de precios del cotizador.
Usa recuadros (tablas con bordes) como espacios para llenar.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:color'), '888888')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def add_field_box(doc, label, height_cm=1.2, placeholder="RD$ _______"):
    """Label arriba, caja abajo para llenar."""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.height = Cm(height_cm)
    set_cell_border(cell)
    set_cell_bg(cell, 'FFFAEB')
    cell_p = cell.paragraphs[0]
    cell_run = cell_p.add_run(placeholder)
    cell_run.font.size = Pt(10)
    cell_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_pricing_row_table(doc, headers, rows_data, col_widths=None):
    """Tabla de precios con headers y filas vacías para llenar."""
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.style = 'Table Grid'

    # Headers
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, '713DFF')
        cell_p = cell.paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Rows
    for r_idx, row in enumerate(rows_data, start=1):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_border(cell)
            if c_idx == 0:
                set_cell_bg(cell, 'F5F5F5')
            else:
                set_cell_bg(cell, 'FFFAEB')
            cell_p = cell.paragraphs[0]
            run = cell_p.add_run(val)
            run.font.size = Pt(10)
            if c_idx == 0:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_section_heading(doc, number, title):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x71, 0x3D, 0xFF)
    p.paragraph_format.space_after = Pt(8)


def generate_pricing_form():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ========== HEADER ==========
    p = doc.add_paragraph()
    run = p.add_run("MEMORAMA — Tabla de Precios para Cotizador Online")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x71, 0x3D, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Llena esta tabla con tus precios reales. Esto alimenta el cotizador automático de tu tienda Shopify.")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== INSTRUCCIONES ==========
    p = doc.add_paragraph()
    run = p.add_run("⚡ Cómo llenar este formato:")
    run.bold = True
    run.font.size = Pt(11)

    instrucciones = [
        "Todos los precios van en RD$ (pesos dominicanos)",
        "Si algo no aplica, escribe 'N/A'",
        "Si tienes variaciones adicionales que no están aquí, agrégalas en la sección 9 (Observaciones)",
        "Los descuentos por volumen aplican al TOTAL del pedido, no al precio unitario"
    ]
    for i in instrucciones:
        doc.add_paragraph(i, style='List Bullet')

    # ========== 1. PRECIO BASE POR PRODUCTO ==========
    add_section_heading(doc, "1", "Precio base por producto")
    p = doc.add_paragraph()
    run = p.add_run("Precio del producto SIN personalización (camiseta, polo, gorra, etc). Si cada producto tiene precio distinto, llena la tabla. Si todos van al mismo precio base, solo llena la primera fila.")
    run.font.size = Pt(10)

    productos_base = [
        ["Camiseta manga corta", "RD$ _______"],
        ["Camiseta manga larga", "RD$ _______"],
        ["Polo / camisa tipo Columbia", "RD$ _______"],
        ["Hoodie / sudadera", "RD$ _______"],
        ["Gorra", "RD$ _______"],
        ["Tote bag / bolso", "RD$ _______"],
        ["Otro: __________", "RD$ _______"],
        ["Otro: __________", "RD$ _______"],
    ]
    add_pricing_row_table(doc, ["Producto", "Precio base (RD$)"], productos_base)

    # ========== 2. TÉCNICAS DE PERSONALIZACIÓN ==========
    add_section_heading(doc, "2", "Técnicas de personalización")
    p = doc.add_paragraph()
    run = p.add_run("Costo ADICIONAL por técnica (se suma al precio base del producto). Ejemplo: si bordado cuesta RD$180 extra y la camiseta base cuesta RD$450, el total con bordado es RD$630.")
    run.font.size = Pt(10)

    tecnicas = [
        ["Bordado (estándar, hasta 10 cm)", "RD$ _______"],
        ["Bordado premium (más de 10 cm)", "RD$ _______"],
        ["DTF (transfer digital)", "RD$ _______"],
        ["Serigrafía 1 color", "RD$ _______"],
        ["Serigrafía 2 colores", "RD$ _______"],
        ["Serigrafía 3 colores", "RD$ _______"],
        ["Serigrafía 4+ colores", "RD$ _______"],
        ["Sublimación full color", "RD$ _______"],
        ["Vinil textil", "RD$ _______"],
        ["Otro: __________", "RD$ _______"],
    ]
    add_pricing_row_table(doc, ["Técnica", "Costo adicional (RD$)"], tecnicas)

    # ========== 3. POSICIONES DE ESTAMPADO ==========
    add_section_heading(doc, "3", "Posiciones de estampado")
    p = doc.add_paragraph()
    run = p.add_run("Costo ADICIONAL por cada posición donde se estampa. Si el cliente pide pecho + espalda, se suman ambos.")
    run.font.size = Pt(10)

    posiciones = [
        ["Pecho izquierdo (pequeño)", "RD$ _______"],
        ["Pecho completo (centro)", "RD$ _______"],
        ["Espalda completa", "RD$ _______"],
        ["Manga izquierda", "RD$ _______"],
        ["Manga derecha", "RD$ _______"],
        ["Cuello trasero (label)", "RD$ _______"],
        ["Otro: __________", "RD$ _______"],
    ]
    add_pricing_row_table(doc, ["Posición", "Costo adicional (RD$)"], posiciones)

    # ========== 4. DESCUENTOS POR VOLUMEN ==========
    add_section_heading(doc, "4", "Descuentos por volumen")
    p = doc.add_paragraph()
    run = p.add_run("Porcentaje de descuento según cantidad de unidades en el pedido. Se aplica sobre el subtotal. Si no das descuento en un tier, escribe 0%.")
    run.font.size = Pt(10)

    volumen = [
        ["1 – 10 unidades", "_____ %"],
        ["11 – 25 unidades", "_____ %"],
        ["26 – 50 unidades", "_____ %"],
        ["51 – 100 unidades", "_____ %"],
        ["101 – 200 unidades", "_____ %"],
        ["201 – 500 unidades", "_____ %"],
        ["500+ unidades", "_____ %"],
    ]
    add_pricing_row_table(doc, ["Cantidad", "% Descuento"], volumen)

    # ========== 5. PEDIDO MÍNIMO ==========
    add_section_heading(doc, "5", "Pedido mínimo")
    add_field_box(doc, "¿Cuál es el pedido mínimo por producto?", placeholder="_____ unidades")
    add_field_box(doc, "¿Hay algún producto que tenga un mínimo distinto? (especifica cuál)", placeholder="Ej: Bordados mínimo 20 un, DTF mínimo 10 un")

    # ========== 6. TIEMPO DE ENTREGA ==========
    add_section_heading(doc, "6", "Tiempo de entrega")
    p = doc.add_paragraph()
    run = p.add_run("Días hábiles desde que se confirma el pedido (con logo aprobado + 50% pago).")
    run.font.size = Pt(10)

    tiempos = [
        ["Bordado", "_____ días"],
        ["DTF", "_____ días"],
        ["Serigrafía", "_____ días"],
        ["Sublimación", "_____ días"],
        ["Vinil", "_____ días"],
    ]
    add_pricing_row_table(doc, ["Técnica", "Tiempo de entrega"], tiempos)

    add_field_box(doc, "¿Ofreces servicio express / urgente? ¿Cuánto cobras extra?", height_cm=1.5, placeholder="Ej: Sí, +30% del total. Tiempo: 3 días")

    # ========== 7. CARGOS ADICIONALES ==========
    add_section_heading(doc, "7", "Cargos adicionales opcionales")

    cargos = [
        ["Digitalización de logo (para bordado)", "RD$ _______"],
        ["Separación de colores (serigrafía)", "RD$ _______"],
        ["Prueba física antes de producir (sample)", "RD$ _______"],
        ["Empaque individual", "RD$ _______"],
        ["Envío a domicilio (Santo Domingo)", "RD$ _______"],
        ["Envío a domicilio (otras provincias)", "RD$ _______"],
    ]
    add_pricing_row_table(doc, ["Concepto", "Costo (RD$)"], cargos)

    # ========== 8. POLÍTICA DE PAGO ==========
    add_section_heading(doc, "8", "Política de pago")
    add_field_box(doc, "¿Qué % de anticipo pides para empezar?", placeholder="_____ %")
    add_field_box(doc, "¿Métodos de pago aceptados?", height_cm=1.5, placeholder="Ej: Transferencia, PayPal, efectivo, tarjeta")
    add_field_box(doc, "¿Das factura fiscal?", placeholder="Sí / No")

    # ========== 9. OBSERVACIONES ==========
    add_section_heading(doc, "9", "Observaciones / variaciones adicionales")
    p = doc.add_paragraph()
    run = p.add_run("Cualquier cosa que afecte el precio y no esté arriba: tipos de hilo, marcas de tela, tipos de tinta especial, acabados, etc.")
    run.font.size = Pt(10)

    add_field_box(doc, "Notas del cliente", height_cm=4, placeholder="Escribe aquí cualquier detalle importante...")

    # ========== FOOTER ==========
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("—" * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Preparado por MM Agency · Contacto: +1 849-577-2978")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar
    output_path = "/Users/martinmercedes/Desktop/Executive assistant 2/outputs/briefings/memorama-tabla-precios-cotizador.docx"
    doc.save(output_path)
    print(f"✅ Generado: {output_path}")


if __name__ == "__main__":
    generate_pricing_form()
