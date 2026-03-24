#!/usr/bin/env python3
"""
Genera un cuestionario/briefing de cliente para gestión de contenido en redes sociales.
El documento llenado sirve como contexto completo para generación de contenido con IA.

Uso:
    python3 generar_briefing.py                    # Solo template genérico
    python3 generar_briefing.py --nombre anabel    # Template + copia personalizada
"""

import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = Path(__file__).resolve().parent.parent.parent
TEMPLATE_DIR = BASE_DIR / "templates" / "briefing-cliente"
OUTPUT_DIR = BASE_DIR / "outputs" / "briefings"

# Colores
DARK_GRAY = RGBColor(0x2D, 0x2D, 0x2D)
MEDIUM_GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
ACCENT_COLOR = RGBColor(0x1A, 0x1A, 0x2E)  # Azul oscuro profesional
TABLE_HEADER_BG = "1A1A2E"
TABLE_ALT_BG = "F5F5F8"


def setup_styles(doc):
    """Configura estilos del documento."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1 - Título principal
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = ACCENT_COLOR
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(4)

    # Heading 2 - Secciones
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = ACCENT_COLOR
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 - Subsecciones
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = DARK_GRAY
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_separator(doc):
    """Agrega una línea separadora sutil."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_header(doc, client_name=None):
    """Agrega el header del documento."""
    title = "Briefing de Cliente"
    subtitle = "Cuestionario para Gestión de Contenido en Redes Sociales"

    doc.add_heading(title, level=1)

    if client_name:
        p = doc.add_paragraph()
        run = p.add_run(f"Cliente: {client_name.title()}")
        run.font.size = Pt(14)
        run.font.color.rgb = MEDIUM_GRAY
        run.font.italic = True

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = MEDIUM_GRAY

    add_separator(doc)

    # Instrucciones
    p = doc.add_paragraph()
    run = p.add_run("Instrucciones: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "Completa cada sección con la mayor cantidad de detalle posible. "
        "No hay respuestas incorrectas — mientras más información nos des, "
        "mejor será el contenido que creamos para ti. "
        "Tiempo estimado: 25-30 minutos."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = MEDIUM_GRAY
    p.paragraph_format.space_after = Pt(16)


def add_question_field(doc, number, question, rows=1, hint=None):
    """Agrega una pregunta con campo de respuesta (tabla)."""
    # Pregunta
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {question}")
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    if hint:
        p_hint = doc.add_paragraph()
        run = p_hint.add_run(hint)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = LIGHT_GRAY
        p_hint.paragraph_format.space_after = Pt(4)

    # Campo de respuesta
    table = doc.add_table(rows=rows, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        row.height = Cm(1.2) if rows == 1 else Cm(0.8)
        cell = row.cells[0]
        cell.width = Cm(16)
        # Placeholder text
        p = cell.paragraphs[0]
        run = p.add_run("")
        run.font.size = Pt(10)
        set_cell_shading(cell, "FAFAFA")

    doc.add_paragraph()  # Espaciado


def add_list_question(doc, number, question, items, hint=None):
    """Agrega una pregunta con tabla de múltiples filas para listar items."""
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {question}")
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    if hint:
        p_hint = doc.add_paragraph()
        run = p_hint.add_run(hint)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = LIGHT_GRAY
        p_hint.paragraph_format.space_after = Pt(4)

    # Tabla con headers
    table = doc.add_table(rows=1 + len(items), cols=len(items[0]) if items else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(items[0] if items else []):
        cell = header_row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for row_idx in range(1, len(items)):
        row = table.rows[row_idx]
        row.height = Cm(0.9)
        for col_idx, cell_text in enumerate(items[row_idx]):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            if row_idx % 2 == 0:
                set_cell_shading(cell, TABLE_ALT_BG)

    doc.add_paragraph()


def add_checkbox_table(doc, number, question, options, hint=None):
    """Agrega una pregunta con opciones tipo checkbox."""
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {question}")
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

    if hint:
        p_hint = doc.add_paragraph()
        run = p_hint.add_run(hint)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = LIGHT_GRAY
        p_hint.paragraph_format.space_after = Pt(4)

    # Tabla: Marca | Pilar | Descripción | Notas
    headers = ["✓", "Pilar", "Descripción", "Notas"]
    table = doc.add_table(rows=1 + len(options), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(7)
        row.cells[3].width = Cm(4)

    # Header
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Options
    for idx, (pilar, desc) in enumerate(options):
        row = table.rows[idx + 1]
        row.height = Cm(0.8)

        # Checkbox column (empty for client to mark)
        cell_check = row.cells[0]
        p = cell_check.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("☐")
        run.font.size = Pt(14)

        # Pilar name
        cell_name = row.cells[1]
        p = cell_name.paragraphs[0]
        run = p.add_run(pilar)
        run.font.size = Pt(10)
        run.bold = True

        # Description
        cell_desc = row.cells[2]
        p = cell_desc.paragraphs[0]
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.color.rgb = MEDIUM_GRAY

        # Notes (empty)
        cell_notes = row.cells[3]
        p = cell_notes.paragraphs[0]
        run = p.add_run("")

        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_ALT_BG)

    doc.add_paragraph()


def add_section(doc, number, title):
    """Agrega un heading de sección."""
    doc.add_heading(f"Sección {number} — {title}", level=2)


def build_document(client_name=None):
    """Construye el documento completo."""
    doc = Document()

    # Page setup: A4, margins 2.5cm
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_header(doc, client_name)

    # ─── SECCIÓN 1: DATOS GENERALES ───
    add_section(doc, 1, "Datos Generales del Negocio")
    add_question_field(doc, "1.1", "Nombre de la empresa o marca")
    add_question_field(doc, "1.2", "Nombre del dueño o persona principal")
    add_question_field(doc, "1.3", "Industria o nicho",
                       hint="Ej: estética, fitness, restaurante, consultoría, salud...")
    add_question_field(doc, "1.4", "Ciudad y país donde operan")
    add_question_field(doc, "1.5", "Sitio web (si tienen)")
    add_question_field(doc, "1.6", "Redes sociales actuales (incluye los links)",
                       rows=3)
    add_question_field(doc, "1.7", "¿Cuánto tiempo llevan operando?")

    # ─── SECCIÓN 2: PRODUCTOS Y SERVICIOS ───
    add_section(doc, 2, "Productos y Servicios")
    add_question_field(doc, "2.1",
                       "Lista de productos o servicios principales",
                       rows=5,
                       hint="Escribe el nombre y una descripción breve de cada uno.")
    add_question_field(doc, "2.2",
                       "¿Cuál es el producto o servicio estrella? (el que más vende o el que quieren impulsar)")
    add_question_field(doc, "2.3", "Rango de precios",
                       hint="Ej: desde $500 hasta $5,000 RD")
    add_question_field(doc, "2.4",
                       "¿Qué hace diferente tu producto o servicio vs. la competencia?",
                       rows=2,
                       hint="Tu propuesta de valor — por qué alguien te elegiría a ti.")
    add_question_field(doc, "2.5",
                       "Promociones o combos que manejan frecuentemente",
                       rows=2)

    # ─── SECCIÓN 3: CLIENTE IDEAL ───
    add_section(doc, 3, "Cliente Ideal (Audiencia)")
    add_question_field(doc, "3.1", "Rango de edad del cliente ideal",
                       hint="Ej: 25-40 años")
    add_question_field(doc, "3.2", "Género predominante",
                       hint="Mujeres, hombres, ambos...")
    add_question_field(doc, "3.3", "¿Dónde vive tu cliente ideal?",
                       hint="Ciudad, zona, país")
    add_question_field(doc, "3.4", "Nivel socioeconómico aproximado",
                       hint="Bajo, medio, medio-alto, alto")
    add_question_field(doc, "3.5",
                       "¿Qué problema o necesidad resuelve tu producto/servicio para ellos?",
                       rows=2)
    add_question_field(doc, "3.6",
                       "¿Qué frase diría tu cliente ideal cuando busca lo que tú ofreces?",
                       rows=2,
                       hint="Ej: \"Necesito bajar de peso pero no tengo tiempo\", \"Quiero verse más joven sin cirugía\"")
    add_question_field(doc, "3.7",
                       "¿Qué objeciones tiene tu cliente antes de comprar?",
                       rows=2,
                       hint="Ej: \"Es muy caro\", \"No sé si funciona\", \"No tengo tiempo\"")
    add_question_field(doc, "3.8",
                       "¿Dónde pasa más tiempo tu cliente en redes sociales?",
                       hint="Instagram, TikTok, Facebook, YouTube, WhatsApp...")

    # ─── SECCIÓN 4: IDENTIDAD DE MARCA Y VOZ ───
    add_section(doc, 4, "Identidad de Marca y Voz")
    add_question_field(doc, "4.1",
                       "Si tu marca fuera una persona, ¿cómo la describirías en 3 palabras?",
                       hint="Ej: cercana, experta, divertida / elegante, seria, confiable")
    add_question_field(doc, "4.2",
                       "Tono de comunicación (elige 1 o 2)",
                       hint="Opciones: formal, casual, inspiracional, técnico, humorístico, motivacional")
    add_question_field(doc, "4.3",
                       "¿Tuteas o hablas de usted a tu audiencia?")
    add_question_field(doc, "4.4",
                       "Palabras o frases que usas mucho en tu negocio",
                       rows=2,
                       hint="Vocabulario propio de tu marca, muletillas, expresiones que te identifican.")
    add_question_field(doc, "4.5",
                       "Palabras o temas PROHIBIDOS",
                       rows=2,
                       hint="Cosas que nunca quieres que se mencionen en tu contenido.")
    add_question_field(doc, "4.6",
                       "¿Tienes un slogan o tagline?",
                       hint="Si no tienes, déjalo en blanco.")

    # ─── SECCIÓN 5: ESTILO VISUAL ───
    add_section(doc, 5, "Estilo Visual")
    add_question_field(doc, "5.1",
                       "Colores de la marca",
                       hint="Si tienes códigos hex, inclúyelos. Si no, describe los colores (ej: azul marino, dorado).")
    add_question_field(doc, "5.2",
                       "Tipografía o fuentes que usan (si las conocen)")
    add_question_field(doc, "5.3",
                       "Estilo visual general",
                       hint="Minimalista, colorido, elegante, rústico, moderno, tropical, limpio...")
    add_question_field(doc, "5.4",
                       "3 cuentas de Instagram cuyo estilo visual admiras (no tienen que ser competidores)",
                       rows=2,
                       hint="Incluye los @usernames")
    add_question_field(doc, "5.5",
                       "¿Tienen fotos profesionales del negocio, productos o equipo?",
                       hint="Sí/No — y cómo nos las comparten (Google Drive, WhatsApp, etc.)")
    add_question_field(doc, "5.6",
                       "¿Usan algún preset o filtro específico para fotos?")

    # ─── SECCIÓN 6: COMPETENCIA ───
    add_section(doc, 6, "Competencia")
    add_question_field(doc, "6.1",
                       "Nombra 2-3 competidores directos (con link a sus redes si es posible)",
                       rows=3)
    add_question_field(doc, "6.2",
                       "¿Qué hacen bien ellos?",
                       rows=2)
    add_question_field(doc, "6.3",
                       "¿Qué hacen mal o qué oportunidad ves vs. ellos?",
                       rows=2)
    add_question_field(doc, "6.4",
                       "¿Por qué un cliente te elegiría a TI en vez de a ellos?",
                       rows=2)

    # ─── SECCIÓN 7: PILARES DE CONTENIDO ───
    add_section(doc, 7, "Pilares de Contenido")
    add_checkbox_table(doc, "7",
                       "Marca los pilares que aplican para tu marca y agrega notas si quieres:",
                       [
                           ("Educativo", "Tips, datos, tutoriales, información útil"),
                           ("Producto/Servicio", "Mostrar lo que vendes, demos, antes/después"),
                           ("Testimonios", "Casos de éxito, reseñas, resultados de clientes"),
                           ("Detrás de cámaras", "Día a día, proceso de trabajo, equipo"),
                           ("Entretenimiento", "Tendencias, memes, contenido viral adaptado"),
                           ("Promociones", "Ofertas, descuentos, combos especiales"),
                           ("Otro:", "(Escribe aquí)"),
                           ("Otro:", "(Escribe aquí)"),
                       ],
                       hint="Los pilares son las categorías de temas sobre los que siempre vamos a crear contenido.")

    # ─── SECCIÓN 8: OBJETIVOS Y METAS ───
    add_section(doc, 8, "Objetivos y Metas")
    add_question_field(doc, "8.1",
                       "Objetivo principal en redes sociales",
                       hint="Más seguidores, más ventas, más leads/consultas, posicionamiento de marca, otro...")
    add_question_field(doc, "8.2",
                       "Meta concreta para los próximos 3 meses",
                       hint="Ej: \"Conseguir 50 leads\", \"Vender 20 unidades\", \"Llegar a 5,000 seguidores\"")
    add_question_field(doc, "8.3",
                       "¿Tienen presupuesto para pauta publicitaria (Meta Ads)?",
                       hint="Sí/No — si sí, ¿cuánto mensual aproximadamente?")
    add_question_field(doc, "8.4",
                       "¿Qué métricas les importan más?",
                       hint="Alcance, engagement, leads, ventas, tráfico web, seguidores...")

    # ─── SECCIÓN 9: LOGÍSTICA DE CONTENIDO ───
    add_section(doc, 9, "Logística de Contenido")
    add_question_field(doc, "9.1",
                       "Frecuencia deseada de publicación",
                       hint="Ej: 3 posts + 5 stories por semana, 4 reels al mes...")
    add_question_field(doc, "9.2",
                       "Formatos preferidos",
                       hint="Posts estáticos, carruseles, reels, stories, lives — marca todos los que apliquen")
    add_question_field(doc, "9.3",
                       "¿Quién aparecerá en cámara para reels o videos?",
                       hint="Nombre y rol. Si nadie, indicarlo.")
    add_question_field(doc, "9.4",
                       "¿Quién aprueba el contenido antes de publicar?",
                       hint="Nombre y canal de comunicación preferido para enviar las revisiones.")
    add_question_field(doc, "9.5",
                       "Canal de comunicación preferido para revisiones",
                       hint="WhatsApp, email, otro...")
    add_question_field(doc, "9.6",
                       "Fechas especiales o temporadas importantes para el negocio",
                       rows=3,
                       hint="Ej: San Valentín, Black Friday, aniversario del negocio, temporada alta...")

    # ─── SECCIÓN 10: INFORMACIÓN ADICIONAL ───
    add_section(doc, 10, "Información Adicional")
    add_question_field(doc, "10.1",
                       "¿Hay algo que NO quieres que hagamos con tu contenido?",
                       rows=2)
    add_question_field(doc, "10.2",
                       "¿Hay algo que sí o sí quieres que incluyamos?",
                       rows=2)
    add_question_field(doc, "10.3",
                       "Referencia de un post o reel que te haya gustado mucho (link)",
                       rows=2)
    add_question_field(doc, "10.4",
                       "Notas adicionales — cualquier cosa que creas relevante",
                       rows=3)

    # Footer
    add_separator(doc)
    p = doc.add_paragraph()
    run = p.add_run("CONFIDENCIAL — Este documento contiene información privada del cliente.")
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GRAY
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    parser = argparse.ArgumentParser(
        description="Genera un briefing de cliente para gestión de contenido en redes sociales."
    )
    parser.add_argument(
        "--nombre",
        type=str,
        default=None,
        help="Nombre del cliente (para personalizar el documento)."
    )
    args = parser.parse_args()

    # Asegurar directorios
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Generar template genérico
    doc_template = build_document()
    template_path = TEMPLATE_DIR / "briefing-cliente-template.docx"
    doc_template.save(str(template_path))
    print(f"✓ Template guardado en: {template_path}")

    # Generar copia personalizada si se da nombre
    if args.nombre:
        doc_named = build_document(client_name=args.nombre)
        nombre_clean = args.nombre.lower().replace(" ", "-")
        named_path = OUTPUT_DIR / f"briefing-{nombre_clean}.docx"
        doc_named.save(str(named_path))
        print(f"✓ Briefing personalizado guardado en: {named_path}")


if __name__ == "__main__":
    main()
