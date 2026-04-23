"""
Generate Casa 360 briefing as Word (.docx) with MM Agency branding
Estilo con RECUADROS visibles para llenar (como formulario profesional)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand colors MM Agency
YELLOW = RGBColor(0xFA, 0xFA, 0x00)
VIOLET = RGBColor(0x71, 0x3D, 0xFF)
DARK = RGBColor(0x18, 0x18, 0x18)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
LIGHT_GREY = RGBColor(0x9A, 0x9A, 0x9A)
BORDER_GREY = RGBColor(0xCC, 0xCC, 0xCC)
BOX_BG = RGBColor(0xFA, 0xFA, 0xFA)
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def set_cell_border(cell, color="CCCCCC", size="6"):
    """Set all borders of a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_bg(cell, color="FAFAFA"):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def add_field_box(label, height_lines=1, bg="FAFAFA"):
    """Add a labeled field with a bordered box for filling"""
    # Label
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    # Box (1x1 table)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.5)

    # Empty lines inside box for filling
    for i in range(height_lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(' ')
        run.font.size = Pt(11)

    set_cell_border(cell, color="CCCCCC", size="6")
    set_cell_bg(cell, color=bg)

    # Spacing after
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_small_field(label, height_lines=1):
    """Smaller field for short answers"""
    add_field_box(label, height_lines=height_lines)


def add_large_field(label, height_lines=4):
    """Larger field for long answers"""
    add_field_box(label, height_lines=height_lines)


def add_two_column_fields(label1, label2):
    """Two fields side by side"""
    # Labels row
    labels_table = doc.add_table(rows=1, cols=2)
    labels_table.autofit = False

    cell1 = labels_table.cell(0, 0)
    cell2 = labels_table.cell(0, 1)
    cell1.width = Cm(8.1)
    cell2.width = Cm(8.1)

    p1 = cell1.paragraphs[0]
    run = p1.add_run(label1)
    run.font.size = Pt(9)
    run.bold = True

    p2 = cell2.paragraphs[0]
    run = p2.add_run(label2)
    run.font.size = Pt(9)
    run.bold = True

    # Boxes row
    boxes_table = doc.add_table(rows=1, cols=2)
    boxes_table.autofit = False

    b1 = boxes_table.cell(0, 0)
    b2 = boxes_table.cell(0, 1)
    b1.width = Cm(8.1)
    b2.width = Cm(8.1)

    b1.paragraphs[0].add_run(' ')
    b2.paragraphs[0].add_run(' ')

    set_cell_border(b1, color="CCCCCC", size="6")
    set_cell_border(b2, color="CCCCCC", size="6")
    set_cell_bg(b1, color="FAFAFA")
    set_cell_bg(b2, color="FAFAFA")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading_bar(num, text, color=DARK):
    """Section heading with number and colored bar"""
    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Number
    p = doc.add_paragraph()
    run = p.add_run(num)
    run.font.size = Pt(10)
    run.font.color.rgb = VIOLET
    run.bold = True
    p.paragraph_format.space_after = Pt(2)

    # Title
    p2 = doc.add_paragraph()
    run = p2.add_run(text.upper())
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = DARK
    p2.paragraph_format.space_after = Pt(4)

    # Bar under title (narrow yellow line)
    p3 = doc.add_paragraph()
    run = p3.add_run('▬▬▬▬▬')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFA, 0xC8, 0x00)
    p3.paragraph_format.space_after = Pt(10)


def add_checkbox_list(items, columns=1):
    """Add checkbox list, optionally in columns"""
    if columns == 1:
        for item in items:
            p = doc.add_paragraph()
            run = p.add_run('☐  ')
            run.font.size = Pt(12)
            run = p.add_run(item)
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.space_after = Pt(3)
    else:
        # Organize in table
        rows_needed = (len(items) + columns - 1) // columns
        table = doc.add_table(rows=rows_needed, cols=columns)
        table.autofit = True
        for idx, item in enumerate(items):
            row = idx // columns
            col = idx % columns
            cell = table.cell(row, col)
            p = cell.paragraphs[0]
            run = p.add_run('☐  ')
            run.font.size = Pt(11)
            run = p.add_run(item)
            run.font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_intro_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(8)


def add_subsection_label(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GREY
    run.bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_divider_space():
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ============================================
# HEADER + COVER
# ============================================

# Small brand mark at top
header_p = doc.add_paragraph()
run = header_p.add_run('MM AGENCY.')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = DARK
header_p.paragraph_format.space_after = Pt(0)

sub = doc.add_paragraph()
run = sub.add_run('DOCUMENTO DE BRIEFING  ·  20 / ABR / 2026  ·  REF: MM-2026-008')
run.font.size = Pt(7)
run.font.color.rgb = LIGHT_GREY
sub.paragraph_format.space_after = Pt(24)

# Big title
title = doc.add_paragraph()
run = title.add_run('BRIEFING')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = DARK
title.paragraph_format.space_after = Pt(0)

subtitle = doc.add_paragraph()
run = subtitle.add_run('CASA 360')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = VIOLET
subtitle.paragraph_format.space_after = Pt(0)

tagline = doc.add_paragraph()
run = tagline.add_run('Ebanistería  ·  República Dominicana')
run.font.size = Pt(12)
run.italic = True
run.font.color.rgb = GREY
tagline.paragraph_format.space_after = Pt(18)

add_intro_text('Completa este documento llenando los recuadros. Puedes escribir directamente en Word, abrirlo desde el celular, o imprimirlo y llenarlo a mano. Entre más detalle compartas, mejor podremos conectar tu marca con los clientes correctos.')

# ============================================
# 01. NEGOCIO
# ============================================
add_heading_bar('01', 'Información del negocio')

add_two_column_fields('Nombre comercial oficial', 'Nombre legal / RNC')
add_two_column_fields('Ubicación / zona donde opera', 'Años en el negocio')
add_two_column_fields('Dueño / contacto principal', 'Teléfono / WhatsApp del negocio')
add_two_column_fields('Email del negocio', 'Tamaño del taller / equipo')

# ============================================
# 02. SERVICIOS
# ============================================
add_heading_bar('02', 'Servicios de ebanistería')
add_intro_text('Marca los servicios que ofrece Casa 360:')

add_checkbox_list([
    'Cocinas modulares / integrales',
    'Clósets y walk-in closets',
    'Muebles de baño / vanities',
    'Puertas de madera a medida',
    'Escaleras de madera',
    'Libreros / bibliotecas',
    'Centros de entretenimiento / TV',
    'Muebles de oficina a medida',
    'Restauración de muebles antiguos',
    'Proyectos comerciales (locales, oficinas, hoteles)',
    'Otro:',
], columns=2)

add_large_field('¿Cuál servicio deja más dinero? (el que queremos impulsar)', height_lines=2)
add_large_field('¿Cuál servicio tiene más demanda pero margen menor?', height_lines=2)
add_large_field('¿Algún servicio "estrella" que diferencie a Casa 360?', height_lines=2)

# ============================================
# 03. MADERAS Y ACABADOS
# ============================================
add_heading_bar('03', 'Tipos de madera y acabados')

add_large_field('Maderas con las que trabajan (caoba, roble, pino, cedro, melamina, MDF, otros)', height_lines=2)
add_large_field('Acabados disponibles (barniz, laca, tinte, natural, pintura)', height_lines=2)
add_two_column_fields('¿Herrajes importados o locales?', '¿Fabrican + instalan, o solo fabrican?')

# ============================================
# 04. TIPO DE PROYECTO
# ============================================
add_heading_bar('04', 'Tipo de proyecto')

add_subsection_label('Rango de precio típico de un proyecto')
add_two_column_fields('Precio mínimo (RD$)', 'Precio promedio (RD$)')
add_small_field('Precio máximo (RD$)', height_lines=1)

add_two_column_fields('Tiempo promedio de un proyecto', '¿Cuánto piden de anticipo?')
add_small_field('¿Dan garantía? ¿De cuánto tiempo?', height_lines=1)

# ============================================
# 05. CLIENTE IDEAL
# ============================================
add_heading_bar('05', 'Cliente ideal')

add_subsection_label('Perfil del cliente que más le conviene')
add_two_column_fields('Edad aproximada', 'Nivel socioeconómico')
add_small_field('Zonas de RD donde más venden', height_lines=1)

add_large_field('¿Cliente final (casa propia) o profesionales (arquitectos, diseñadores, constructoras)?', height_lines=2)
add_large_field('¿De dónde vienen sus clientes hoy? (referidos, Google, Instagram, Facebook, boca a boca)', height_lines=3)
add_large_field('¿Qué busca ese cliente cuando llega a Casa 360? (calidad, precio, diseño, tiempo, acabados)', height_lines=3)
add_large_field('¿Qué miedos o dudas tiene ese cliente antes de contratar?', height_lines=3)

# ============================================
# 06. DIFERENCIADORES
# ============================================
add_heading_bar('06', 'Diferenciadores')

add_large_field('¿Por qué un cliente debería elegir Casa 360 y no otra ebanistería?', height_lines=4)
add_large_field('¿Qué hacen mejor que la competencia? (acabados, diseño, tiempos, precio, atención, materiales)', height_lines=4)
add_large_field('¿Tienen proyectos emblemáticos o clientes conocidos que puedan servir como prueba social?', height_lines=3)
add_large_field('¿Cuáles son los trabajos más impresionantes que han hecho? (para destacar en la página y contenido)', height_lines=4)

# ============================================
# 07. COMPETENCIA
# ============================================
add_heading_bar('07', 'Competencia local')

add_large_field('Competidores directos en su zona / segmento (nombres, Instagram, páginas web)', height_lines=4)
add_large_field('¿Qué hacen bien ellos?', height_lines=3)
add_large_field('¿Qué hacen mal o podrían mejorar?', height_lines=3)
add_small_field('Precio promedio de la competencia vs. Casa 360', height_lines=1)

# ============================================
# 08. IDENTIDAD VISUAL
# ============================================
add_heading_bar('08', 'Identidad visual')

add_subsection_label('¿Tiene logo?')
add_checkbox_list([
    'Sí — pasar archivo en alta calidad',
    'No — hay que diseñarlo',
])

add_small_field('¿Tiene colores definidos de marca? ¿Cuáles?', height_lines=2)

add_subsection_label('Estilo visual que quiere transmitir')
add_checkbox_list([
    'Moderno / minimalista',
    'Lujo / premium (gama alta)',
    'Artesanal / cálido (madera natural, hecho a mano)',
    'Industrial / contemporáneo',
    'Profesional / corporativo',
    'Otro:',
], columns=2)

add_subsection_label('Redes sociales existentes')
add_two_column_fields('Instagram', 'Facebook')
add_two_column_fields('TikTok', 'YouTube')

add_subsection_label('¿Tiene banco de fotos/videos de proyectos?')
add_checkbox_list([
    'Sí — cantidad aproximada:',
    'Algunos pero no suficientes',
    'No — hay que producirlos',
])

# ============================================
# 09. OBJETIVOS
# ============================================
add_heading_bar('09', 'Objetivos del negocio')

add_large_field('¿Qué quiere lograr en los próximos 3 meses con la página + contenido? (ej: 10 cotizaciones al mes, posicionarse en [zona], captar clientes de gama alta)', height_lines=4)
add_two_column_fields('¿Cuántos proyectos al mes le gustaría cerrar?', 'Meta de facturación mensual ideal')
add_large_field('¿Quiere crecer en volumen o en ticket promedio? (más clientes o clientes que paguen más)', height_lines=2)

# ============================================
# 10. OPERACIÓN
# ============================================
add_heading_bar('10', 'Operación y atención al cliente')

add_subsection_label('¿Cómo atiende clientes nuevos hoy?')
add_checkbox_list([
    'WhatsApp',
    'Llamada telefónica',
    'Visita al taller',
    'Visita al cliente para medir',
], columns=2)

add_large_field('¿Tiene un proceso de cotización definido? (cómo pasa de "interesado" a "cliente")', height_lines=3)

add_subsection_label('¿Quién responde los mensajes en redes?')
add_checkbox_list([
    'Él mismo',
    'Asistente / secretaria',
    'Nadie sistemáticamente',
])

add_two_column_fields('Horario de atención', '¿Tiene showroom o solo taller?')

# ============================================
# 11. CONTENIDO
# ============================================
add_heading_bar('11', 'Contenido para redes')

add_large_field('¿Puede documentar su trabajo diario? (fotos/videos del taller, proceso, entregas)', height_lines=2)

add_subsection_label('¿Está dispuesto a grabar videos cortos mostrando?')
add_checkbox_list([
    'Cortes de madera',
    'Armado de piezas',
    'Acabados',
    'Antes y después',
    'Entregas / instalación',
], columns=2)

add_large_field('¿Tiene testimonios de clientes satisfechos? (videos, reseñas, fotos)', height_lines=2)
add_small_field('¿Un miembro del equipo puede pasar fotos/videos a WhatsApp cada semana?', height_lines=1)

# ============================================
# 12. WEB Y EMBUDO
# ============================================
add_heading_bar('12', 'Sitio web y embudo')

add_subsection_label('Dominio')
add_checkbox_list([
    'Ya tiene dominio',
    'Hay que comprarlo',
])
add_small_field('Nombre de dominio preferido (casa360.com.do / casa360rd.com / ebanisteriacasa360.com)', height_lines=1)

add_subsection_label('¿Qué quiere que haga la página principal?')
add_checkbox_list([
    'Cotización online',
    'Galería de proyectos',
    'Contacto rápido por WhatsApp',
    'Blog / contenido educativo',
    'Testimonios',
], columns=2)

add_large_field('Objetivo del embudo (ej: captar el teléfono → llamarle para cotizar → cerrar venta)', height_lines=3)

# ============================================
# 13. CRM
# ============================================
add_heading_bar('13', 'CRM (GoHighLevel)')

add_small_field('¿Quién manejará el CRM día a día?', height_lines=1)

add_subsection_label('¿Qué quiere trackear?')
add_checkbox_list([
    'Leads entrantes',
    'Cotizaciones enviadas',
    'Proyectos en curso',
    'Clientes pagados',
    'Referidos',
], columns=2)

add_subsection_label('Automatizaciones que quiere')
add_checkbox_list([
    'Mensaje automático cuando llega un lead',
    'Recordatorio de seguimiento',
    'Pedir reseña al entregar proyecto',
    'Otro:',
])

# ============================================
# 14. ACCESOS
# ============================================
add_heading_bar('14', 'Accesos técnicos')
add_intro_text('Los necesitaremos después del pago inicial para poder arrancar:')

add_checkbox_list([
    'Logo en alta calidad (si tiene)',
    'Fotos/videos de proyectos anteriores (mínimo 20 para arrancar)',
    'Acceso a Instagram (contraseña o añadir admin)',
    'Acceso a Facebook Business',
    'WhatsApp Business configurado',
    'Dominio comprado o decidir nombre',
    'Email profesional (crear si no tiene)',
])

# ============================================
# 15. OBSERVACIONES
# ============================================
add_heading_bar('15', 'Observaciones finales')

add_large_field('Cualquier cosa importante que quiera compartir:', height_lines=8)

# ============================================
# PRÓXIMOS PASOS
# ============================================
add_heading_bar('→', 'Próximos pasos', color=VIOLET)

steps = [
    ('1', 'Completar este briefing (WhatsApp o llamada de 30 min)'),
    ('2', 'Enviar materiales (logo, fotos, videos, accesos)'),
    ('3', 'Confirmar pago inicial:'),
]

for num, step in steps:
    p = doc.add_paragraph()
    run = p.add_run(f'  {num}.  ')
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = VIOLET
    run = p.add_run(step)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

# Payment sub-items
payments = [
    'US$500  —  inicial Casa 360 (al entregar Memorama)',
    'US$97  —  GHL mes 1',
    'RD$18,500  —  contenido mes 1 (ambas marcas)',
]
for pay in payments:
    p = doc.add_paragraph()
    run = p.add_run(f'         ·  {pay}')
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(2)

# Step 4
p = doc.add_paragraph()
run = p.add_run('  4.  ')
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = VIOLET
run = p.add_run('Arrancar producción: web + embudo + primer lote de contenido')
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(12)

# Dates
add_two_column_fields('Fecha objetivo de entrega web', 'Fecha de primer post publicado')

# ============================================
# FIRMA
# ============================================
add_heading_bar('✓', 'Aceptación', color=VIOLET)
add_intro_text('Al firmar abajo, aceptas los términos y montos indicados arriba.')

add_two_column_fields('Cliente — Firma y nombre', 'MM Agency — Martin Mercedes')

# ============================================
# FOOTER
# ============================================
doc.add_paragraph().paragraph_format.space_after = Pt(12)

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('MM AGENCY.  ·  hola@mmagency.com.do  ·  mmagency.com.do')
run.font.size = Pt(8)
run.font.color.rgb = LIGHT_GREY
run.italic = True

# Save
output_path = '/Users/martinmercedes/Desktop/Executive assistant 2/outputs/briefings/briefing-casa-360.docx'
import os
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Briefing saved to: {output_path}")
